"""Extracción de contenido desde archivos de entrada.

Código copiado desde main. No contiene lógica tomada de mao-dev-branch.
"""

import io
import pandas as pd


def extract_txt(uploaded_file):
    raw = uploaded_file.getvalue()
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("No fue posible decodificar el archivo de texto.")


def extract_pdf(uploaded_file):
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError(
            "La dependencia 'pypdf' no está disponible en el entorno. "
            "Verifica requirements.txt y vuelve a desplegar la aplicación."
        ) from exc

    reader = PdfReader(io.BytesIO(uploaded_file.getvalue()))
    pages = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        pages.append(f"[Página {page_number}]\n{text}")

    content = "\n\n".join(pages).strip()
    if not content:
        raise ValueError(
            "El PDF se abrió correctamente, pero no contiene texto extraíble. "
            "Si es un PDF escaneado, esta versión requiere OCR."
        )
    return content


def extract_docx(uploaded_file):
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "La dependencia 'python-docx' no está disponible. "
            "Verifica requirements.txt y vuelve a desplegar la aplicación."
        ) from exc

    doc = Document(io.BytesIO(uploaded_file.getvalue()))
    blocks = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            blocks.append(" | ".join(cell.text.strip() for cell in row.cells))

    content = "\n".join(blocks).strip()
    if not content:
        raise ValueError("El DOCX se abrió correctamente, pero no contiene texto.")
    return content


def extract_xlsx(uploaded_file):
    """Extract visible sheet content from an Excel reference file."""
    data = uploaded_file.getvalue()
    sheets = pd.read_excel(io.BytesIO(data), sheet_name=None, dtype=str)
    parts = []
    for sheet_name, df in sheets.items():
        df = df.fillna("")
        parts.append(f"=== HOJA: {sheet_name} ===")
        parts.append(df.to_csv(index=False))
    return "\n".join(parts)


def extract_csv(uploaded_file):
    data = uploaded_file.getvalue()
    try:
        df = pd.read_csv(io.BytesIO(data), dtype=str)
    except Exception:
        df = pd.read_csv(io.BytesIO(data), dtype=str, encoding="latin-1")
    return df.fillna("").to_csv(index=False)


def extract_source(uploaded_file):
    extension = uploaded_file.name.rsplit(".", 1)[-1].lower()

    if extension in ("txt", "md"):
        content = extract_txt(uploaded_file)
    elif extension == "pdf":
        content = extract_pdf(uploaded_file)
    elif extension == "docx":
        content = extract_docx(uploaded_file)
    elif extension in {"xlsx", "xls"}:
        content = extract_xlsx(uploaded_file)
    elif extension == "csv":
        content = extract_csv(uploaded_file)
    else:
        raise ValueError(f"Formato no soportado: .{extension}")

    if not content.strip():
        raise ValueError("El documento no contiene contenido utilizable.")

    return content
