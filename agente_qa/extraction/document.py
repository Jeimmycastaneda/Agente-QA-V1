"""Extracción de PDF, DOCX, TXT, XLSX y CSV."""

import io
import pandas as pd

def extract_txt(uploaded_file):
    raw = uploaded_file.getvalue()
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            pass
    raise ValueError("No fue posible decodificar el archivo de texto.")

def extract_pdf(uploaded_file):
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(uploaded_file.getvalue()))
    pages = []
    for n, page in enumerate(reader.pages, 1):
        pages.append(f"[Página {n}]\n{page.extract_text() or ''}")
    content = "\n\n".join(pages).strip()
    if not content:
        raise ValueError("El PDF no contiene texto extraíble. Si es escaneado, requiere OCR.")
    return content

def extract_docx(uploaded_file):
    from docx import Document
    doc = Document(io.BytesIO(uploaded_file.getvalue()))
    blocks = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            blocks.append(" | ".join(cell.text.strip() for cell in row.cells))
    content = "\n".join(blocks).strip()
    if not content:
        raise ValueError("El DOCX no contiene texto.")
    return content

def extract_xlsx(uploaded_file):
    sheets = pd.read_excel(io.BytesIO(uploaded_file.getvalue()), sheet_name=None, dtype=str)
    parts = []
    for name, df in sheets.items():
        parts.append(f"=== HOJA: {name} ===")
        parts.append(df.fillna("").to_csv(index=False))
    return "\n".join(parts)

def extract_csv(uploaded_file):
    data = uploaded_file.getvalue()
    try:
        df = pd.read_csv(io.BytesIO(data), dtype=str)
    except Exception:
        df = pd.read_csv(io.BytesIO(data), dtype=str, encoding="latin-1")
    return df.fillna("").to_csv(index=False)

def extract_source(uploaded_file):
    ext = uploaded_file.name.rsplit(".", 1)[-1].lower()
    handlers = {
        "txt": extract_txt, "md": extract_txt, "pdf": extract_pdf,
        "docx": extract_docx, "xlsx": extract_xlsx, "xls": extract_xlsx,
        "csv": extract_csv,
    }
    if ext not in handlers:
        raise ValueError(f"Formato no soportado: .{ext}")
    content = handlers[ext](uploaded_file)
    if not content.strip():
        raise ValueError("El documento no contiene contenido utilizable.")
    return content
