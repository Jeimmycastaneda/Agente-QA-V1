import io
import json
import logging
import os
import re
import time
from pathlib import Path
from datetime import datetime
import html
from html import escape

import pandas as pd
import streamlit as st
from editor_azure import render_azure_style_editor, delete_test_case

import base64
from urllib.error import HTTPError, URLError
from urllib.parse import quote, urlencode
from urllib.request import Request, urlopen

class AzureDevOpsError(RuntimeError):
    pass

def _az_config():
    try:
        secrets = st.secrets
    except Exception:
        secrets = {}
    return {
        "org": str(secrets.get("AZURE_DEVOPS_ORG", os.getenv("AZURE_DEVOPS_ORG", ""))).strip(),
        "project": str(secrets.get("AZURE_DEVOPS_PROJECT", os.getenv("AZURE_DEVOPS_PROJECT", ""))).strip(),
        "pat": str(secrets.get("AZURE_DEVOPS_PAT", os.getenv("AZURE_DEVOPS_PAT", ""))).strip(),
    }

def _az_validate(cfg):
    missing = [k for k in ("org", "project", "pat") if not cfg.get(k)]
    if missing:
        raise AzureDevOpsError("Faltan secretos de Azure DevOps: " + ", ".join(missing))

def _az_get_json(url, pat):
    token = base64.b64encode(f":{pat}".encode("utf-8")).decode("ascii")
    req = Request(
        url,
        method="GET",
        headers={
            "Authorization": f"Basic {token}",
            "Accept": "application/json",
            "User-Agent": "Agente-QA-Streamlit/1.0",
        },
    )
    try:
        with urlopen(req, timeout=20) as response:
            return json.loads(response.read().decode("utf-8")), response.headers
    except HTTPError as exc:
        detail = ""
        try:
            detail = exc.read().decode("utf-8", errors="replace")[:1000]
        except Exception:
            pass
        if exc.code in (401, 403):
            raise AzureDevOpsError(
                f"Azure rechazó la autenticación/autorización (HTTP {exc.code}). "
                "Verifica el PAT, su vigencia y sus scopes."
            ) from exc
        if exc.code == 404:
            raise AzureDevOpsError(
                "No se encontró el proyecto o el recurso de Test Plans. "
                "Verifica AZURE_DEVOPS_ORG y AZURE_DEVOPS_PROJECT."
            ) from exc
        raise AzureDevOpsError(f"Azure respondió HTTP {exc.code}. {detail}") from exc
    except URLError as exc:
        raise AzureDevOpsError(
            f"No fue posible comunicarse con Azure DevOps: {exc.reason}"
        ) from exc

def test_connection():
    cfg = _az_config()
    _az_validate(cfg)
    org = quote(cfg["org"], safe="")
    project = quote(cfg["project"], safe="")
    url = (
        f"https://dev.azure.com/{org}/{project}/_apis/wit/workitemtypes/"
        f"Test%20Case?api-version=7.1"
    )
    payload, _ = _az_get_json(url, cfg["pat"])
    return {
        "organization": cfg["org"],
        "project": cfg["project"],
        "work_item_type": payload.get("name", "Test Case"),
        "message": "Conexión correcta. Solo se realizó una consulta de lectura.",
    }

def _az_testplan_url(cfg, path):
    return (
        f"https://dev.azure.com/{quote(cfg['org'], safe='')}/"
        f"{quote(cfg['project'], safe='')}/_apis/testplan/{path}"
    )

def list_test_plans(limit=10):
    # SOLO GET. Trae únicamente los 10 Test Plans más recientes por ID.
    cfg = _az_config()
    _az_validate(cfg)
    params = urlencode({"api-version": "7.1", "$top": int(limit)})
    url = _az_testplan_url(cfg, "plans") + "?" + params
    payload, _ = _az_get_json(url, cfg["pat"])
    plans = payload.get("value") or []
    plans = sorted(plans, key=lambda x: int(x.get("id") or 0), reverse=True)[:int(limit)]
    rows = []
    for p in plans:
        rows.append({
            "id": p.get("id"),
            "name": _ui_text(p.get("name"), "Sin nombre"),
            "state": p.get("state", ""),
            "area_path": p.get("areaPath", ""),
            "iteration": p.get("iteration", ""),
            "start_date": p.get("startDate", ""),
            "end_date": p.get("endDate", ""),
        })
    return {
        "ok": True,
        "organization": cfg["org"],
        "project": cfg["project"],
        "count": len(rows),
        "plans": rows,
        "message": (
            "Consulta correcta. Solo se consultaron los 10 Test Plans más recientes "
            "por ID; no se creó, modificó ni eliminó ningún recurso."
        ),
    }

def list_test_suites(plan_id):
    cfg = _az_config()
    _az_validate(cfg)
    path = f"Plans/{quote(str(plan_id), safe='')}/suites"
    payload, _ = _az_get_json(
        _az_testplan_url(cfg, path) + "?api-version=7.1",
        cfg["pat"],
    )
    rows = []
    for s in payload.get("value") or []:
        plan = s.get("plan") if isinstance(s.get("plan"), dict) else {}
        parent = s.get("parentSuite") if isinstance(s.get("parentSuite"), dict) else {}
        rows.append({
            "id": s.get("id"),
            "name": _ui_text(s.get("name"), "Suite sin nombre"),
            "suite_type": s.get("suiteType", ""),
            "plan_id": plan.get("id", plan_id),
            "parent_suite": parent.get("id"),
        })
    return rows

def list_test_cases(plan_id, suite_id):
    cfg = _az_config()
    _az_validate(cfg)
    path = (
        f"Plans/{quote(str(plan_id), safe='')}/Suites/"
        f"{quote(str(suite_id), safe='')}/TestCase"
    )
    payload, _ = _az_get_json(
        _az_testplan_url(cfg, path) + "?api-version=7.1",
        cfg["pat"],
    )

    rows = []
    for item in payload.get("value") or []:
        # Azure puede devolver el Work Item en testCase, workItem o directamente.
        wi = {}
        for key in ("testCase", "workItem"):
            candidate = item.get(key)
            if isinstance(candidate, dict):
                wi = candidate
                break
        if not wi:
            wi = item if isinstance(item, dict) else {}

        # ID robusto: primero id directo; luego posibles referencias/URLs.
        # IMPORTANTE: no usar una expresión condicional sin paréntesis aquí.
        # Azure normalmente devuelve el ID directamente en testCase.id.
        # La versión anterior podía convertir ese ID válido en None
        # porque el "else None" afectaba toda la expresión.
        nested_work_item = wi.get("workItem")
        nested_id = (
            nested_work_item.get("id")
            if isinstance(nested_work_item, dict)
            else None
        )
        raw_id = (
            wi.get("id")
            or item.get("id")
            or nested_id
        )
        if not raw_id:
            for container in (wi, item):
                for key in ("url", "href"):
                    value = container.get(key) if isinstance(container, dict) else None
                    if value:
                        match = re.search(r"/workitems/(\d+)(?:[/?]|$)", str(value), re.I)
                        if match:
                            raw_id = match.group(1)
                            break
                if raw_id:
                    break

        title = (
            wi.get("name")
            or wi.get("title")
            or item.get("name")
            or item.get("title")
        )

        # No mostrar registros sin ID: son referencias incompletas y producen "None —".
        if raw_id is None or str(raw_id).strip() == "":
            continue

        rows.append({
            "id": str(raw_id).strip(),
            "title": _ui_text(title, "Test Case sin título"),
            "raw": item,
        })

    return rows

def _az_parse_steps_xml(xml_text):
    if not xml_text:
        return []
    decoded = html.unescape(str(xml_text))
    steps = []
    for match in re.finditer(r"<step\b[^>]*>.*?</step>", decoded, flags=re.I | re.S):
        node = match.group(0)
        vals = re.findall(
            r"<parameterizedString[^>]*>(.*?)</parameterizedString>",
            node,
            flags=re.I | re.S,
        )
        clean = []
        for value in vals[:2]:
            value = re.sub(r"<[^>]+>", "", value)
            clean.append(html.unescape(value).strip())
        if clean:
            steps.append({
                "Step #": len(steps) + 1,
                "Action": clean[0] if len(clean) > 0 else "",
                "Expected value": clean[1] if len(clean) > 1 else "",
            })
    return steps

def get_test_case_detail(test_case_id):
    cfg = _az_config()
    _az_validate(cfg)
    org = quote(cfg["org"], safe="")
    project = quote(cfg["project"], safe="")
    url = (
        f"https://dev.azure.com/{org}/{project}/_apis/wit/workitems/"
        f"{quote(str(test_case_id), safe='')}?api-version=7.1"
    )
    payload, _ = _az_get_json(url, cfg["pat"])
    fields = payload.get("fields") or {}
    return {
        "id": payload.get("id"),
        "title": fields.get("System.Title", ""),
        "description": fields.get("System.Description", "") or "",
        "steps": _az_parse_steps_xml(fields.get("Microsoft.VSTS.TCM.Steps", "") or ""),
        "area_path": fields.get("System.AreaPath", ""),
        "iteration_path": fields.get("System.IterationPath", ""),
        "state": fields.get("System.State", ""),
        "work_item_type": fields.get("System.WorkItemType", "Test Case"),
        "raw_fields": fields,
    }

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import HRFlowable, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


# ===== V30: COBERTURA OBLIGATORIA POR CU =====
def _normalize_cu(value):
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value).strip()).casefold()


def _extract_related_cu(tc):
    """Extrae un único CU desde Related Use Case, aceptando:
    CU-324, CU-324 - Nombre, CU-324: Nombre, etc.
    """
    value = (
        tc.get("Related Use Case")
        or tc.get("related_use_case")
        or tc.get("use_case")
        or tc.get("Requirement / Use Case")
        or tc.get("Caso de uso relacionado")
        or ""
    )
    if isinstance(value, list):
        raw_parts = [str(x).strip() for x in value if str(x).strip()]
    else:
        raw_parts = [x.strip() for x in re.split(r"[;\n|]", str(value)) if x.strip()]

    results = []
    for part in raw_parts:
        # El ID del CU es la referencia trazable. Acepta CU-324 y variantes
        # con separadores/nombre después del ID.
        match = re.search(r"\b(CU[-_ ]?\d+)\b", part, flags=re.IGNORECASE)
        if match:
            results.append(match.group(1).upper().replace("_", "-").replace(" ", "-"))
        else:
            results.append(part)
    return results


def _use_case_catalog(identified_use_cases):
    """Normaliza el catálogo de CU identificado en la HU/documentación."""
    catalog = []
    for cu in identified_use_cases or []:
        if isinstance(cu, dict):
            cid = safe_text(
                cu.get("ID"), cu.get("id"), cu.get("Use Case ID"), cu.get("CU")
            )
            name = safe_text(
                cu.get("Name"), cu.get("name"), cu.get("Title"), cu.get("Description")
            )
        else:
            cid = safe_text(cu)
            name = cid
        if cid or name:
            catalog.append({"id": cid, "name": name or cid})
    return catalog


def _resolve_related_use_cases(data, source_content=""):
    """Resuelve Related Use Case sin inventar, siguiendo la prioridad aprobada.

    1. Usa Related Use Case si ya existe.
    2. Si falta, busca evidencia en 'Casos de Uso Relacionados'.
    3. Si aún falta, usa el título/nombre del CU identificado en la HU.
    4. Si no hay una coincidencia única, deja vacío y genera alerta.
    """
    cases = data.get("TEST_CASES", []) or []
    catalog = _use_case_catalog(data.get("USE_CASES", []))
    if not catalog:
        return

    source = safe_text(source_content)
    related_block = ""
    if source:
        marker = re.search(r"Casos de Uso Relacionados\s*:?(.*?)(?=\n\s*(?:Precondiciones|Criterios de Aceptación|Descripción|Resultado esperado|Historia de Usuario|Casos de Uso)\b|\Z)", source, flags=re.I | re.S)
        if marker:
            related_block = marker.group(1)

    related_norm = _normalize_cu(related_block)

    for tc in cases:
        current = safe_text(
            tc.get("Related Use Case"),
            tc.get("RelatedUseCase"),
            tc.get("Caso de uso relacionado"),
        )
        if current and _extract_related_cu(tc):
            continue

        searchable = " ".join(
            safe_text(tc.get(k))
            for k in (
                "Title", "Scenario", "Description", "Criterion",
                "Requirement / Use Case", "use_case", "related_use_case"
            )
        )
        searchable_norm = _normalize_cu(searchable)

        # 2) Buscar primero en la sección explícita de Casos de Uso Relacionados.
        matches = []
        for cu in catalog:
            cid = _normalize_cu(cu["id"])
            name = _normalize_cu(cu["name"])
            if (cid and cid in related_norm) or (name and name in related_norm):
                matches.append(cu)

        # Si el bloque tiene varias relaciones, usar la que además coincida con
        # el escenario del CP; nunca elegir arbitrariamente entre varias.
        if len(matches) > 1:
            contextual = [
                cu for cu in matches
                if (_normalize_cu(cu["id"]) and _normalize_cu(cu["id"]) in searchable_norm)
                or (_normalize_cu(cu["name"]) and _normalize_cu(cu["name"]) in searchable_norm)
            ]
            if len(contextual) == 1:
                matches = contextual

        if len(matches) == 1:
            tc["Related Use Case"] = f"{matches[0]['id']} - {matches[0]['name']}"
            continue

        # 3) Si no existe relación explícita, buscar el título/nombre del CU
        # dentro del escenario/Descripción generado a partir de la HU.
        title_matches = []
        for cu in catalog:
            name_norm = _normalize_cu(cu["name"])
            cid_norm = _normalize_cu(cu["id"])
            if (name_norm and name_norm in searchable_norm) or (cid_norm and cid_norm in searchable_norm):
                title_matches.append(cu)

        if len(title_matches) == 1:
            tc["Related Use Case"] = f"{title_matches[0]['id']} - {title_matches[0]['name']}"
            continue

        # 4) No inventar. La generación queda trazable y el dato se valida antes
        # de crear cualquier CP en Azure.
        tc["Related Use Case"] = ""
        alerts = tc.get("Alerts")
        if not isinstance(alerts, list):
            alerts = []
            tc["Alerts"] = alerts
        alerts.append({
            "Alert": "Caso de uso relacionado no identificado de forma inequívoca",
            "Reason": "No se encontró una relación explícita ni una coincidencia única con los Casos de Uso identificados en la Historia de Usuario.",
            "Validation Required": "Validar el Caso de Uso relacionado con el equipo funcional antes de crear el CP en Azure.",
        })


def calculate_cu_coverage(cases, identified_use_cases):
    """Mínimo 1 CP por cada CU y exactamente 1 CU por CP."""
    cu_map = {}
    for cu in identified_use_cases or []:
        if isinstance(cu, dict):
            cid = str(
                cu.get("ID") or cu.get("id") or
                cu.get("Use Case ID") or cu.get("CU") or ""
            ).strip()
            name = str(
                cu.get("Name") or cu.get("name") or
                cu.get("Title") or cu.get("Description") or ""
            ).strip()
        else:
            cid = str(cu).strip()
            name = cid
        if cid:
            cu_map[_normalize_cu(cid)] = {"id": cid, "name": name or cid}

    covered = {}
    cp_without_cu = []
    cp_multiple_cu = []

    for index, tc in enumerate(cases or [], start=1):
        cp_id = str(tc.get("ID") or f"CP-{index:05d}").strip()
        relations = _extract_related_cu(tc)

        if len(relations) == 0:
            cp_without_cu.append(cp_id)
            continue
        if len(relations) != 1:
            cp_multiple_cu.append(cp_id)
            continue

        rel = _normalize_cu(relations[0])
        matched = None
        for cu_key, cu_info in cu_map.items():
            if rel == cu_key or rel == _normalize_cu(cu_info["name"]):
                matched = cu_key
                break

        if matched is None:
            cp_without_cu.append(cp_id)
        else:
            covered.setdefault(matched, []).append(cp_id)

    missing = [
        info for key, info in cu_map.items()
        if key not in covered
    ]

    total_cu = len(cu_map)
    total_cp = len(cases or [])
    covered_count = len(covered)
    percentage = round((covered_count / total_cu) * 100, 1) if total_cu else 0.0

    return {
        "total_cu": total_cu,
        "total_cp": total_cp,
        "covered_cu": covered_count,
        "missing_cu": missing,
        "cp_without_cu": cp_without_cu,
        "cp_multiple_cu": cp_multiple_cu,
        "percentage": percentage,
        "valid": (
            total_cu > 0
            and not missing
            and not cp_without_cu
            and not cp_multiple_cu
        ),
    }


def validate_minimum_cu_coverage(data):
    """Bloquea cualquier resultado que no cubra todos los CU."""
    cases = data.get("TEST_CASES", []) or []
    use_cases = data.get("USE_CASES", []) or []

    if not use_cases:
        raise ValueError(
            "GENERACIÓN BLOQUEADA: Gemini no devolvió la lista completa de Casos de Uso (USE_CASES)."
        )

    metrics = calculate_cu_coverage(cases, use_cases)

    if not metrics["valid"]:
        missing = ", ".join(
            f'{x["id"]} - {x["name"]}'
            for x in metrics["missing_cu"]
        )
        details = []
        if missing:
            details.append("CU sin CP: " + missing)
        if metrics["cp_without_cu"]:
            details.append(
                "CP sin CU válido: " + ", ".join(metrics["cp_without_cu"])
            )
        if metrics["cp_multiple_cu"]:
            details.append(
                "CP con más de un CU: " + ", ".join(metrics["cp_multiple_cu"])
            )
        raise ValueError(
            f'COBERTURA INCOMPLETA: {metrics["covered_cu"]}/'
            f'{metrics["total_cu"]} CU cubiertos '
            f'({metrics["percentage"]}%). '
            + " | ".join(details)
        )

    return metrics


def render_cu_coverage(metrics):
    st.markdown("### 📊 Cobertura mínima por Caso de Uso")
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("CU identificados", metrics["total_cu"])
    c2.metric("CP generados", metrics["total_cp"])
    c3.metric("CU cubiertos", metrics["covered_cu"])
    c4.metric("Cobertura", f'{metrics["percentage"]}%')

    if metrics["valid"]:
        st.success("✅ Cobertura completa: cada CU tiene mínimo un CP.")
    else:
        st.error(
            f'🔴 Cobertura incompleta: {metrics["covered_cu"]}/'
            f'{metrics["total_cu"]} CU cubiertos. '
            f'Faltan {len(metrics["missing_cu"])} CU.'
        )
        if metrics["missing_cu"]:
            with st.expander("Ver CU sin Caso de Prueba", expanded=True):
                for cu in metrics["missing_cu"]:
                    st.write(f'• **{cu["id"]}** — {cu["name"]}')
        if metrics["cp_without_cu"]:
            st.warning(
                "CP sin CU válido: " + ", ".join(metrics["cp_without_cu"])
            )
        if metrics["cp_multiple_cu"]:
            st.warning(
                "CP relacionados con más de un CU: "
                + ", ".join(metrics["cp_multiple_cu"])
            )

try:
    from google import genai
    from google.genai import types
except ImportError:
    genai = None
    types = None

APP_VERSION = "V36-ULTIMOS10-REFERENCIA-Y-REVISION-FUNCIONAL"
MODEL = "gemini-3.6-flash"
FALLBACK_MODELS = [
    "gemini-3.6-flash",
    "gemini-3.5-flash-lite",
]

# ============================================================
# COLUMNAS APROBADAS — NO AGREGAR NI CAMBIAR TÍTULOS
# ============================================================
# Columnas EXACTAS requeridas por Azure DevOps Test Plans para importación XLSX.
# Para casos nuevos, ID queda vacío. Cada step es una fila y repite los campos del CP y establece Tipo Origen Proyecto = Proyecto.
AZURE_COLUMNS = [
    "ID", "Work Item Type", "Title", "Description", "Test Step", "Step Action",
    "Step Expected", "Area Path", "IDPadre", "Tipo Origen Proyecto",
    "Tiempo Real", "Assigned To", "State"
]

MATRIZ_COLUMNS = [
    "TestCaseId", "Title", "Requirement / Use Case", "Criterion", "Scenario",
    "Scenario Type", "Description", "Preconditions", "Validation Method",
    "Coverage", "Alerts", "Effort"
]

EXCEL_CONFIGS = {
    "Autos Colectivos": {
        "sheet_name": "Azure Import",
        "base_id": 10001,
        "base_testpoint": 1001,
        "configuration": "Default configuration",
        "tester": "",
        "title_prefix": "CP-AC-",
        "user_default": "Usuario registrado",
        "steps_with_users": True,
        "area_path": "COTIZADORES WEB\\DESARROLLO",
        "assigned_to": "",
    },
    "Siniestros Fasecolda": {
        "sheet_name": "28443;Fase 3 - RENK170 Siniestr",
        "base_id": 28454,
        "base_testpoint": 6552,
        "configuration": "Default configuration created @ 26/05/2023 15:22:17",
        "tester": "Isabel Cristina Mejía López",
        "title_prefix": "CP-ACSF-",
        "user_default": "Suscriptor Oficina Principal, suscriptor sucursal autos",
        "steps_with_users": True,
        "area_path": "COTIZADORES WEB\\DESARROLLO",
        "assigned_to": "",
    },
    "General QA": {
        "sheet_name": "Casos de Prueba",
        "base_id": 20001,
        "base_testpoint": 2001,
        "configuration": "",
        "tester": "",
        "title_prefix": "CP-",
        "user_default": "",
        "steps_with_users": False,
        "area_path": "COTIZADORES WEB\\DESARROLLO",
        "assigned_to": "",
    },
}

SCHEMA = {
    "type": "object",
    "properties": {
        "USE_CASES": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "ID": {"type": "string"},
                    "Name": {"type": "string"}
                },
                "required": ["ID", "Name"]
            }
        },
        "TEST_CASES": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "ID": {"type": "string"},
                    "Title": {"type": "string"},
                    "Description": {"type": "string"},
                    "Expected Result": {"type": "string"},
                    "Preconditions": {"type": "string"},
                    "Product": {"type": "string"},
                    "Module": {"type": "string"},
                    "Related Use Case": {"type": "string"},
                    "Criterion": {"type": "string"},
                    "Scenario": {"type": "string"},
                    "Scenario Type": {"type": "string"},
                    "Effort": {"type": "string"},
                    "Coverage": {"type": "string"},
                    "Validation Method": {"type": "string"},
                    "Steps": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "Step #": {"type": "integer"},
                                "Action": {"type": "string"},
                                "Expected value": {"type": "string"},
                            },
                            "required": ["Step #", "Action", "Expected value"],
                        },
                    },
                    "Alerts": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "Alert": {"type": "string"},
                                "Reason": {"type": "string"},
                                "Validation Required": {"type": "string"},
                            },
                            "required": ["Alert", "Reason", "Validation Required"],
                        },
                    },
                },
                "required": ["ID", "Title", "Description", "Preconditions", "Steps"],
            },
        },
        "ALERTS": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "Alert": {"type": "string"},
                    "Reason": {"type": "string"},
                    "Validation Required": {"type": "string"},
                },
                "required": ["Alert", "Reason", "Validation Required"],
            },
        },
        "COVERAGE": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "Requirement / Use Case": {"type": "string"},
                    "Criterion": {"type": "string"},
                    "Scenario": {"type": "string"},
                    "Test Case": {"type": "string"},
                    "Validation Method": {"type": "string"},
                    "Coverage": {"type": "string"},
                    "Alerts": {"type": "string"},
                },
                "required": ["Requirement / Use Case", "Criterion", "Scenario", "Test Case"],
            },
        },
    },
    "required": ["USE_CASES", "TEST_CASES", "ALERTS", "COVERAGE"],
}


def _ui_text(value, default=""):
    """Texto seguro para UI: nunca muestra el literal None."""
    if value is None:
        return default
    value = str(value).strip()
    return value if value else default


def safe_text(value, default="", *fallbacks):
    """Convierte valores a texto y permite una cadena de valores de respaldo.

    Se mantiene compatible con el uso anterior de safe_text(value, default),
    y además permite safe_text(valor1, valor2, valor3, ...), tomando el
    primer valor no vacío. Esto evita errores cuando la generación combina
    campos alternativos de la fuente.
    """
    candidates = (value, default, *fallbacks)
    for candidate in candidates:
        if candidate is None:
            continue
        if isinstance(candidate, (dict, list)):
            text = json.dumps(candidate, ensure_ascii=False).strip()
        else:
            text = str(candidate).strip()
        if text:
            return text
    return ""


def safe_steps(tc):
    steps = tc.get("Steps", [])
    return steps if isinstance(steps, list) else []


def normalize_coverage(value):
    v = safe_text(value).strip().lower()
    allowed = {
        "completa": "Completa",
        "parcial": "Parcial",
        "no cubierta": "No cubierta",
        "fuera de alcance": "Fuera de alcance",
    }
    return allowed.get(v, safe_text(value, "Pendiente"))


def normalize_validation_method(value):
    v = safe_text(value).strip().lower()
    mapping = {
        "ui": "UI",
        "interfaz": "UI",
        "interfaz de usuario": "UI",
        "bd": "BD",
        "base de datos": "BD",
        "database": "BD",
        "api": "API",
        "web service": "API",
        "web services": "API",
        "mixta": "Mixta",
        "mixto": "Mixta",
    }
    return mapping.get(v, safe_text(value, "Pendiente"))



def build_azure_description(product, module, description, expected, preconditions, related_use_case):
    """Construye la estructura aprobada de Description para Azure sin inventar datos."""
    desc = safe_text(description)
    desc = re.sub(r"(?mi)^\s*Descripción:\s*", "", desc, count=1).strip()
    return (
        f"Producto: {safe_text(product, 'Pendiente')}\n\n"
        f"Módulo: {safe_text(module, 'Pendiente')}\n\n"
        f"Descripción: {desc or 'Pendiente'}\n\n"
        f"Resultado esperado de la prueba: {safe_text(expected, 'Pendiente')}\n\n"
        f"Precondiciones: {safe_text(preconditions, 'Pendiente')}\n\n"
        f"Caso de uso relacionado: {safe_text(related_use_case, 'Pendiente')}"
    )


def format_description_for_azure(description):
    """Formatea Description como Markdown legible en Azure DevOps sin cambiar su contenido.

    Solo transforma presentación: separa bloques, resalta etiquetas y normaliza listas.
    No agrega ni elimina información funcional.
    """
    text = safe_text(description).replace("\r\n", "\n").replace("\r", "\n")
    if not text:
        return ""

    # Normaliza espacios/tabs sin destruir saltos de línea.
    text = re.sub(r"[ \t]+", " ", text)

    # Si el modelo dejó las etiquetas pegadas en una sola línea, las separamos.
    labels = [
        "Producto:",
        "Módulo:",
        "Descripción:",
        "Resultado esperado de la prueba:",
        "Precondiciones:",
        "Caso de uso relacionado:",
    ]
    for label in labels:
        text = re.sub(rf"\s*{re.escape(label)}\s*", f"\\n{label} ", text, count=1)

    # Normaliza viñetas comunes provenientes del modelo.
    text = re.sub(r"(?m)^\s*[•●▪◦]\s*", "- ", text)
    text = re.sub(r"(?m)^\s*[o]\s+", "- ", text)
    text = re.sub(r"\n\s*[-–—]\s*", "\n- ", text)

    # Si hay bullets pegados después de una oración, sepáralos.
    text = re.sub(r"\s+(-\s+)", r"\n\1", text)

    # Limpieza de saltos excesivos.
    text = re.sub(r"\n{3,}", "\n\n", text).strip()

    # Resalta únicamente las etiquetas de los seis bloques aprobados.
    for label in labels:
        text = re.sub(rf"(?m)^{re.escape(label)}\s*", f"**{label}** ", text)

    # Deja una separación visual consistente entre bloques principales.
    for label in labels[1:]:
        text = re.sub(rf"\n\*\*{re.escape(label)}\*\*", f"\n\n**{label}**", text)
    text = re.sub(r"^\*\*Producto:\*\*", "**Producto:**", text)

    return text.strip()

def module_token(module, title="", scenario=""):
    raw = safe_text(module) or safe_text(title) or safe_text(scenario) or "GENERAL"
    raw = re.sub(r"[^A-Za-z0-9]+", " ", raw).strip().upper()
    words = raw.split()
    if not words:
        return "GENERAL"
    if len(words) == 1:
        return words[0][:12]
    return "".join(w[0] for w in words)[:8]


def build_case_title(tc, case_id):
    """
    V12: garantiza que Title sea un título funcional y no el CP ID.
    Prioridad:
      1) Title generado por el modelo si no es solo el ID.
      2) Scenario
      3) Description
      4) Related Use Case
      5) fallback controlado.
    """
    raw_title = safe_text(tc.get("Title"))
    normalized_title = re.sub(r"\s+", " ", raw_title).strip()

    # Un título igual al ID no es un título funcional.
    if (
        not normalized_title
        or normalized_title.upper() == case_id.upper()
        or re.fullmatch(r"CP-[A-Z0-9_-]+-\d{5}", normalized_title.upper())
    ):
        candidates = [
            safe_text(tc.get("Scenario")),
            safe_text(tc.get("Description")),
            safe_text(tc.get("Related Use Case")),
        ]

        for candidate in candidates:
            candidate = re.sub(r"\s+", " ", candidate).strip()
            if candidate and candidate.upper() != case_id.upper():
                normalized_title = candidate
                break

    if not normalized_title:
        normalized_title = f"Caso de prueba {case_id}"

    return normalized_title


def normalize_case_id(raw_id, module, index, prefix="CP-AC-"):
    candidate = safe_text(raw_id)
    if re.fullmatch(r"CP-AC-[A-Za-z0-9_-]+-\d{5}", candidate):
        return candidate
    return f"{prefix}{module_token(module)}-{index:05d}"


def find_coverage(data, tc):
    tc_id = safe_text(tc.get("ID"))
    for row in data.get("COVERAGE", []) if isinstance(data.get("COVERAGE", []), list) else []:
        if safe_text(row.get("Test Case")) == tc_id:
            return row
    return {}


def aggregate_case_alerts(data, tc):
    parts = []
    case_alerts = tc.get("Alerts", [])
    if isinstance(case_alerts, list):
        for alert in case_alerts:
            text = safe_text(alert.get("Alert"))
            reason = safe_text(alert.get("Reason"))
            validation = safe_text(alert.get("Validation Required"))
            if text:
                if reason:
                    text += f": {reason}"
                if validation:
                    text += f" | Validación: {validation}"
                parts.append(text)
    return " | ".join(parts) if parts else "Sin Alertas"


# ============================================================
# EXTRACCIÓN ROBUSTA DE DOCUMENTOS
# ============================================================
def extract_txt(uploaded_file):
    raw = uploaded_file.getvalue()
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("No fue posible decodificar el archivo de texto.")


def extract_pdf(uploaded_file):
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError(
            "La dependencia 'pypdf' no está disponible en el entorno. "
            "Verifica requirements.txt y vuelve a desplegar la aplicación."
        ) from exc

    reader = PdfReader(io.BytesIO(uploaded_file.getvalue()))
    pages = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        pages.append(f"[Página {page_number}]\n{text}")

    content = "\n\n".join(pages).strip()
    if not content:
        raise ValueError(
            "El PDF se abrió correctamente, pero no contiene texto extraíble. "
            "Si es un PDF escaneado, esta versión requiere OCR."
        )
    return content


def extract_docx(uploaded_file):
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "La dependencia 'python-docx' no está disponible. "
            "Verifica requirements.txt y vuelve a desplegar la aplicación."
        ) from exc

    doc = Document(io.BytesIO(uploaded_file.getvalue()))
    blocks = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            blocks.append(" | ".join(cell.text.strip() for cell in row.cells))

    content = "\n".join(blocks).strip()
    if not content:
        raise ValueError("El DOCX se abrió correctamente, pero no contiene texto.")
    return content



def extract_xlsx(uploaded_file):
    """Extract visible sheet content from an Excel reference file."""
    data = uploaded_file.getvalue()
    sheets = pd.read_excel(io.BytesIO(data), sheet_name=None, dtype=str)
    parts = []
    for sheet_name, df in sheets.items():
        df = df.fillna("")
        parts.append(f"=== HOJA: {sheet_name} ===")
        parts.append(df.to_csv(index=False))
    return "\n".join(parts)


def extract_csv(uploaded_file):
    data = uploaded_file.getvalue()
    try:
        df = pd.read_csv(io.BytesIO(data), dtype=str)
    except Exception:
        df = pd.read_csv(io.BytesIO(data), dtype=str, encoding="latin-1")
    return df.fillna("").to_csv(index=False)


def extract_source(uploaded_file):
    extension = uploaded_file.name.rsplit(".", 1)[-1].lower()

    if extension in ("txt", "md"):
        content = extract_txt(uploaded_file)
    elif extension == "pdf":
        content = extract_pdf(uploaded_file)
    elif extension == "docx":
        content = extract_docx(uploaded_file)
    elif extension in {"xlsx", "xls"}:
        content = extract_xlsx(uploaded_file)
    elif extension == "csv":
        content = extract_csv(uploaded_file)
    else:
        raise ValueError(f"Formato no soportado: .{extension}")

    if not content.strip():
        raise ValueError("El documento no contiene contenido utilizable.")

    return content


# ============================================================
# PROMPT EXISTENTE — SE CONSERVA SIN CAMBIOS
# ============================================================
@st.cache_data(ttl=3600)
def load_prompt():
    path = "prompt_qa.txt"
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            content = f.read().strip()
        if content:
            return content

    return (
        "Eres un agente QA especializado en análisis de documentación. "
        "Analiza exclusivamente la fuente proporcionada, no inventes información "
        "y genera TEST_CASES, ALERTS y COVERAGE en JSON."
    )


@st.cache_data(ttl=3600)
def get_valid_models(api_key):
    if genai is None:
        return FALLBACK_MODELS

    try:
        client = genai.Client(api_key=api_key)
        names = []
        for model in client.models.list():
            name = model.name.split("/")[-1]
            if "gemini" in name.lower():
                names.append(name)
        return sorted(set(names)) or FALLBACK_MODELS
    except Exception:
        return FALLBACK_MODELS


def validate_qa_structure(data):
    if not isinstance(data, dict):
        raise ValueError("La respuesta de Gemini no es un objeto JSON.")

    for key in ("USE_CASES", "TEST_CASES", "ALERTS", "COVERAGE"):
        if key not in data:
            raise ValueError(f"Falta la clave requerida: {key}")

    if not isinstance(data["USE_CASES"], list) or not data["USE_CASES"]:
        raise ValueError("Gemini no devolvió los Casos de Uso identificados.")

    if not isinstance(data["TEST_CASES"], list) or not data["TEST_CASES"]:
        raise ValueError("No se generaron casos de prueba.")

    if not isinstance(data["ALERTS"], list):
        data["ALERTS"] = []

    if not isinstance(data["COVERAGE"], list):
        data["COVERAGE"] = []

    return data

def _is_gemini_3x(model_name):
    return safe_text(model_name).lower().startswith(("gemini-3.", "gemini-3"))


def _extract_error_detail(exc):
    raw = str(exc)
    # Keep the useful API message but avoid dumping huge exception payloads.
    return raw[:1800]


def _generate_once(client, model_name, full_prompt):
    """
    V10:
    - Gemini 3.x: no temperature/top_p/top_k.
    - Uses structured JSON output.
    - Limits output to 32K tokens for stability.
    """
    # google-genai actual Python SDK format:
    # response_mime_type + response_schema.
    # Do NOT use response_format here; that shape is rejected by
    # GenerateContentConfig in the installed SDK.
    config = types.GenerateContentConfig(
        response_mime_type="application/json",
        response_schema=SCHEMA,
        max_output_tokens=32768,
    )

    return client.models.generate_content(
        model=model_name,
        contents=full_prompt,
        config=config,
    )


# ============================================================
# ADDENDUM DE CALIDAD — DETALLE FUNCIONAL DEL CP
# ============================================================
DETAILED_QA_ADDENDUM = """
REGLAS OBLIGATORIAS DE NIVEL DE DETALLE PARA LOS CASOS DE PRUEBA

El Caso de Prueba debe reflejar con alto nivel de fidelidad el Caso de Uso (CU)
relacionado. NO generes CP básicos, genéricos ni resumidos.

1. TRAZABILIDAD CU -> CP
- Identifica el CU exacto que sustenta cada CP y usa su contenido completo como
  fuente principal del caso.
- Related Use Case debe indicar el ID y, cuando esté disponible, el nombre del CU.
- Cada CP debe corresponder a EXACTAMENTE un CU.
- Debe existir como mínimo un CP por cada CU identificado.
- Si un CU contiene explícitamente dos o más reglas, condiciones, comportamientos, validaciones o resultados funcionales diferenciados, DEBES generar un CP independiente para cada regla o comportamiento diferenciado. Cada CP debe validar de forma completa la regla que le corresponde, incluyendo contexto, precondiciones, resultado y todos los Steps necesarios.
- Si los requisitos están numerados dentro del CU, analiza cada ítem. Cuando dos ítems representen comportamientos funcionales verificables diferentes, genera CP independientes para cada uno. No agrupes en un mismo CP reglas funcionalmente diferentes solo porque pertenezcan al mismo CU.
- Esta separación NO significa crear un CP por cada Step. Un CP puede contener todos los Steps necesarios para validar completamente su regla.
- Si un requisito está repetido literalmente y no aporta un comportamiento funcional nuevo, no generes un CP adicional por la repetición.
- Si una misma funcionalidad/CU menciona explícitamente dos o más tipos de cotización y cada tipo tiene una regla, cálculo, comportamiento o condición funcional diferenciada, genera CP independientes para cada tipo de cotización y valida cada escenario de forma completa.
- Cada CP independiente debe conservar la trazabilidad al mismo CU cuando corresponda y debe incluir la ruta funcional necesaria para ejecutar su validación específica.

2. DESCRIPCIÓN SUPER DETALLADA
La Description del CP debe explicar el escenario funcional completo. Incluye,
cuando exista en la fuente:
- objetivo y contexto del CU;
- usuario, perfil o rol involucrado;
- módulo, opción, pantalla o funcionalidad;
- condiciones iniciales y precondiciones;
- datos/campos que deben diligenciarse o consultarse;
- reglas de negocio y condiciones;
- estados iniciales/finales;
- restricciones, límites y validaciones;
- comportamiento esperado de cada parte relevante del flujo;
- resultado final que debe obtener el usuario/sistema.

Si para que el CP sea autónomo y ejecutable es necesario incorporar el contenido
completo o casi completo del CU, HAZLO. No reduzcas el CU a una frase genérica.

2A. ESTRUCTURA OBLIGATORIA DE LA DESCRIPTION PARA AZURE DEVOPS
- La columna Description del Excel Azure Import debe contener la descripción funcional completa del CP; NO debe ser una descripción corta ni un resumen del CU.
- La Description debe conservar, cuando exista información en la fuente, el detalle funcional necesario para que el Test Case sea autosuficiente en Azure DevOps.
- La Description debe presentar obligatoriamente, en este orden, estos bloques: Producto, Módulo, Descripción, Resultado esperado de la prueba, Precondiciones y Caso de uso relacionado.
- Producto: usar el producto identificado en la documentación.
- Módulo: usar el módulo/funcionalidad correspondiente al CU.
- Descripción: incluir el contexto funcional completo del escenario, usuarios/perfiles, navegación o funcionalidad, datos/campos, reglas de negocio, condiciones, mensajes, secciones, validaciones, restricciones y cualquier otro detalle explícito relevante del CU.
- Resultado esperado de la prueba: indicar el resultado funcional que debe obtenerse al ejecutar y validar el escenario.
- Precondiciones: conservar las precondiciones explícitas de la fuente y las necesarias para ejecutar el escenario cuando estén sustentadas. No inventar precondiciones.
- Caso de uso relacionado: incluir el ID del CU y su nombre cuando esté disponible.
- Si el CU contiene listas, condiciones, mensajes, nombres de secciones, reglas o referencias a otros CU relevantes para el escenario, conservarlos dentro de Description; no eliminarlos por resumir.
- La Description debe ser específica para el CP y compatible con sus Steps, pero debe conservar el contexto completo necesario del CU.
- No crear una Description genérica como "Validar que el sistema permita...", "Validar la funcionalidad...", "Validar el flujo..." o equivalentes.
- No inventar contenido para completar Producto, Módulo, Resultado esperado, Precondiciones o Caso de uso relacionado. Si un dato no está definido en la fuente, indicarlo como pendiente/por validar según las reglas de no invención.
- La Description debe poder pegarse/importarse directamente en el campo Description de Azure DevOps y conservar saltos de línea, listas y jerarquía funcional cuando sea posible.

3. PASOS COMPLETOS Y EJECUTABLES
- Los Steps deben cubrir TODO el flujo necesario para ejecutar y validar el escenario.
- Cada acción funcional relevante del CU debe aparecer como un paso cuando sea
  necesario para reproducir el escenario.
- Cada paso debe ser concreto y verificable: acción + resultado esperado.
- No agrupes en una sola frase varias acciones importantes si eso hace perder
  trazabilidad o dificulta la ejecución.
- No conviertas cada paso en un CP diferente.
- Un CP puede y debe tener múltiples Steps cuando el CU lo requiera.

4. FIDELIDAD Y NO INVENCIÓN
- Usa exclusivamente la documentación proporcionada como fuente de verdad.
- No inventes usuarios, rutas, URLs, botones, mensajes, campos, valores, reglas,
  permisos, datos o resultados que no estén sustentados.
- Cuando la fuente no defina un dato necesario, conserva la incertidumbre y genera
  ALERTA/Validation Required en lugar de inventarlo.
- Sí puedes reorganizar y redactar para mejorar claridad, pero NO debes perder
  detalles funcionales del CU.

5. CALIDAD MÍNIMA DEL CP
Un CP se considera insuficiente si su Description, Preconditions, Expected Result
 o Steps son tan genéricos que no permiten reconocer qué parte específica del CU
 se está validando.

Ejemplo de calidad insuficiente:
"Validar que el sistema permita realizar la cotización."

Ejemplo de intención correcta:
"Validar el flujo definido en el CU para el perfil indicado, incluyendo el ingreso
al módulo, selección de la opción correspondiente, diligenciamiento/consulta de
los datos definidos, aplicación de las reglas y condiciones establecidas, ejecución
 de la operación y verificación del resultado final especificado por el CU."

El ejemplo anterior solo define el NIVEL DE DETALLE; los datos concretos siempre
deben provenir de la documentación fuente.

6. OPCIONES POSIBLES DE NAVEGACIÓN SIN INVENTAR RUTAS
Cuando el CU no indique de forma explícita la navegación, menú, pantalla u opción exacta,
puedes redactar una **opción posible de acceso o navegación** únicamente si existe
evidencia suficiente en el CU/documentación para plantearla de forma razonable.

REGLAS PARA LA OPCIÓN POSIBLE:
- NO utilizar las etiquetas "Ruta estimada", "Navegación sugerida" ni equivalentes.
- Redactar la opción posible de forma natural dentro de la Description y/o Steps,
  siempre como una alternativa operativa y no como un dato confirmado.
- La opción debe basarse exclusivamente en módulos, funcionalidades, pantallas u
  opciones mencionadas en la documentación.
- No inventar nombres específicos de botones, URLs, IDs, pantallas, menús, rutas o
  secuencias que no tengan sustento en la fuente.
- Si la documentación sí define la navegación, conservarla tal como está sustentada.
- Si no existe evidencia suficiente para proponer una opción de acceso, redactar el
  paso de forma funcional, por ejemplo: "Acceder a la funcionalidad indicada en el CU".
- La opción posible NO debe generar una ruta ficticia ni sustituir el detalle funcional
  del CU. Debe conservar todas las características, condiciones, reglas, mensajes,
  datos, validaciones y resultados definidos en la documentación.

EJEMPLO DE REDACCIÓN:
"Acceder a la funcionalidad indicada en el CU y seleccionar la opción correspondiente
para realizar la operación descrita. La ubicación exacta de la opción debe validarse
con la documentación o el equipo funcional si no está definida en la fuente."

6. EXCEL AZURE — NO CAMBIAR ESTAS REGLAS
- Un CP debe exportarse como un bloque: una fila de cabecera + todas sus filas de Steps.
- En la cabecera: ID, Work Item Type, Title, Area Path, IDPadre, Tipo Origen Proyecto,
  Tiempo Real, Assigned To y State.
- En las filas de pasos: SOLO Test Step, Step Action y Step Expected.
- Tipo Origen Proyecto = Proyecto.
- Area Path = COTIZADORES WEB\\DESARROLLO.
- IDPadre = vacío.
- No crear un CP por cada Step.
"""

def generate_qa_data(
    prompt_text,
    source_content,
    api_key,
    model_name,
    temperature=0.1,
    max_retries=2,
    initial_wait=10,
):
    if genai is None:
        raise RuntimeError("No está instalada la librería google-genai.")

    if not api_key:
        raise ValueError("API Key no configurada.")

    if not source_content.strip():
        raise ValueError("Fuente de información vacía.")

    max_source_chars = 120000
    if len(source_content) > max_source_chars:
        source_content = source_content[:max_source_chars] + (
            "\n...[DOCUMENTO EXCEDE EL LÍMITE DE SEGURIDAD; PRIORIZAR LOS CU Y SU CONTEXTO FUNCIONAL]"
        )

    full_prompt = (
        prompt_text
        + "\n\n==================== ADDENDUM OBLIGATORIO DE CALIDAD ====================\n"
        + DETAILED_QA_ADDENDUM
        + "\n\n==================== FUENTE PROPORCIONADA POR EL USUARIO ====================\n"
        + source_content
        + "\n\n==================== REGLA DE PRIORIDAD ====================\n"
        "La HU/documentación actual es la única fuente de verdad funcional. "
        "Usa el CU completo como fuente principal del CP y conserva sus detalles. "
        "Debe existir mínimo un CP por cada CU y cada CP debe corresponder a un solo CU. "
        "No conviertas Steps en CP. "
        "Related Use Case debe conservar el ID del CU y puede venir como CU-324, CU-324 - nombre o CU-324: nombre; "
        "debe validarse contra los CU reales identificados en USE_CASES. "
        "REGLA OBLIGATORIA PARA CASO DE USO RELACIONADO: si Related Use Case no está definido, busca primero la relación en la sección o campo 'Casos de Uso Relacionados' de la Historia de Usuario/documentación. Si allí no existe una relación utilizable, toma el título/nombre del Caso de Uso que aparece en la Historia de Usuario y úsalo como Caso de uso relacionado, conservando su ID cuando esté disponible. Solo utiliza Casos de Uso presentes en USE_CASES. Si no existe evidencia suficiente o hay más de una coincidencia posible, no inventes: deja el campo vacío y genera ALERTA para validación funcional. "
        "Cuando la navegación no esté definida explícitamente, puedes redactar una opción posible de acceso "
        "solo si existe evidencia suficiente en la documentación. No uses las etiquetas Ruta estimada o "
        "Navegación sugerida, no inventes botones, URLs, menús, pantallas ni rutas, y conserva siempre las "
        "características funcionales y reglas sustentadas en la fuente.\n"
        "\n\n==================== REGLA DE SALIDA ====================\n"
        "Devuelve exclusivamente JSON válido que cumpla el esquema solicitado. "
        "No agregues explicaciones fuera del JSON."
    )

    client = genai.Client(api_key=api_key)

    # Prefer selected model, then use fallbacks.
    # IMPORTANTE: si un modelo devuelve 429/RESOURCE_EXHAUSTED, NO se vuelve a
    # intentar el mismo modelo; se pasa inmediatamente al siguiente candidato.
    candidates = []
    for candidate in [model_name] + FALLBACK_MODELS:
        if candidate and candidate not in candidates:
            candidates.append(candidate)

    errors = []

    for candidate in candidates:
        for attempt in range(max_retries + 1):
            try:
                response = _generate_once(
                    client,
                    candidate,
                    full_prompt,
                )

                response_text = (response.text or "").strip()

                if not response_text:
                    raise RuntimeError(
                        f"{candidate}: Gemini devolvió una respuesta vacía."
                    )

                try:
                    data = json.loads(response_text)
                except json.JSONDecodeError:
                    match = re.search(
                        r"```(?:json)?\s*([\s\S]*?)\s*```",
                        response_text,
                    )
                    if not match:
                        raise RuntimeError(
                            f"{candidate}: la respuesta no es JSON válido. "
                            f"Respuesta: {response_text[:1200]}"
                        )
                    data = json.loads(match.group(1))

                validated = validate_qa_structure(data)
                _resolve_related_use_cases(validated, source_content)
                validate_minimum_cu_coverage(validated)

                st.session_state.quota_exceeded = False
                st.session_state.retry_count = 0

                return validated

            except Exception as exc:
                detail = _extract_error_detail(exc)
                errors.append(f"{candidate} / intento {attempt + 1}: {detail}")

                error_text = detail.lower()

                is_quota = (
                    "429" in detail
                    or "quota" in error_text
                    or "rate limit" in error_text
                    or "resource exhausted" in error_text
                )

                is_retryable_internal = (
                    "500" in detail
                    or "internal" in error_text
                    or "503" in detail
                    or "unavailable" in error_text
                    or "deadline" in error_text
                    or "timeout" in error_text
                )

                if is_quota:
                    st.session_state.quota_exceeded = True
                    st.session_state.retry_count = attempt + 1

                # A 400 caused by an unsupported request/schema should not
                # be retried against the same model; move to next candidate.
                is_bad_request = (
                    "400" in detail
                    or "invalid argument" in error_text
                    or "invalid_argument" in error_text
                    or "unsupported" in error_text
                )

                # 429 / RESOURCE_EXHAUSTED: no insistir sobre el mismo modelo.
                # Saltar inmediatamente al siguiente modelo disponible.
                if is_quota:
                    break

                if (
                    attempt < max_retries
                    and is_retryable_internal
                ):
                    time.sleep(initial_wait * (attempt + 1))
                    continue

                if is_bad_request or is_retryable_internal:
                    break

                # Validation/JSON errors are meaningful and should be shown,
                # but we still allow another model to try.
                break

    joined = "\n\n".join(errors[-8:])
    raise RuntimeError(
        "Gemini no pudo completar la generación con los modelos probados.\n\n"
        "Detalle técnico: " + "\n".join(errors[-8:])
    )


# ============================================================
# EXCEL — ESTRUCTURA APROBADA
# ============================================================
def create_excel(data, config_key):
    """
    Genera:
      1) Hoja 'Azure Import' con la estructura aprobada para Azure, agregando Tipo Origen Proyecto.
      2) Hoja 'Matriz QA' conservando la estructura aprobada.

    Importante:
    - Para CP nuevos, ID queda vacío.
    - Cada CP es un bloque: una fila de cabecera seguida por sus pasos.
    - En las filas de pasos se dejan vacíos los campos de cabecera, igual que en
      el Excel exportado de Azure Test Plans usado como referencia.
    - El ID funcional CP-AC-... se conserva dentro del Title, no en la columna ID.
    """
    config = EXCEL_CONFIGS[config_key]
    output = io.BytesIO()

    azure_rows = []
    matriz_rows = []
    cases = data.get("TEST_CASES", [])

    for idx, tc in enumerate(cases, start=1):
        module = safe_text(tc.get("Module"), "GENERAL")
        case_id = normalize_case_id(
            tc.get("ID"), module, idx, config["title_prefix"]
        )
        title_base = build_case_title(tc, case_id)

        # Conservamos nuestro identificador funcional en el título.
        title = f"{case_id} - {title_base}" if not title_base.startswith(case_id) else title_base

        # Description siempre se exporta con la estructura completa aprobada.
        raw_description = safe_text(tc.get("Description"), safe_text(tc.get("Scenario")))
        preconditions = safe_text(tc.get("Preconditions"))
        scenario = safe_text(tc.get("Scenario"), raw_description)
        description = format_description_for_azure(
            build_azure_description(
                product=safe_text(safe_text(tc.get("Product"), data.get("PRODUCT")), "Pendiente"),
                module=module or "Pendiente",
                description=raw_description,
                expected=safe_text(safe_text(safe_text(tc.get("Expected Result"), tc.get("ExpectedResult")), tc.get("Resultado esperado de la prueba")), "Pendiente"),
                preconditions=preconditions or "Pendiente",
                related_use_case=safe_text(safe_text(safe_text(tc.get("Related Use Case"), tc.get("RelatedUseCase")), tc.get("Caso de uso relacionado")), "Pendiente")
            )
        )
        steps = safe_steps(tc)

        coverage = find_coverage(data, tc)

        validation_method = normalize_validation_method(
            coverage.get(
                "Validation Method",
                tc.get("Validation Method", "Pendiente")
            )
        )

        coverage_value = normalize_coverage(
            coverage.get("Coverage", tc.get("Coverage", "Pendiente"))
        )

        alerts = safe_text(coverage.get("Alerts")) or aggregate_case_alerts(data, tc)

        if alerts == "Sin Alertas" and data.get("ALERTS"):
            general_alerts = []
            for alert in data["ALERTS"]:
                alert_name = safe_text(alert.get("Alert"))
                reason = safe_text(alert.get("Reason"))
                if alert_name:
                    general_alerts.append(
                        f"{alert_name}: {reason}" if reason else alert_name
                    )
            if general_alerts:
                alerts = " | ".join(general_alerts)

        # --------------------------------------------------------
        # AZURE IMPORT — ESTRUCTURA DEL EXPORT DE AZURE TEST PLANS
        # --------------------------------------------------------
        # Un CP = una fila de cabecera + todas sus filas de pasos.
        # La cabecera contiene ID/Work Item Type/Title y metadatos.
        # Las filas siguientes contienen SOLO Test Step/Step Action/Step Expected.
        # Esta estructura evita que Azure interprete cada paso como un nuevo CP.
        area_path = "COTIZADORES WEB\\DESARROLLO"
        assigned_to = safe_text(config.get("assigned_to"))
        state = "Design"
        work_item_type = "Test Case"

        if not steps:
            steps_for_export = [{
                "Step #": 1,
                "Action": "Información insuficiente para definir el paso.",
                "Expected value": "Validar con el equipo funcional antes de ejecutar.",
            }]
        else:
            steps_for_export = steps

        # Fila cabecera: exactamente una por CP.
        azure_rows.append({
            "ID": "",
            "Work Item Type": work_item_type,
            "Title": title,
            "Description": description,
            "Test Step": "",
            "Step Action": "",
            "Step Expected": "",
            "Area Path": area_path,
            "IDPadre": "",
            "Tipo Origen Proyecto": "Proyecto",
            "Tiempo Real": "",
            "Assigned To": assigned_to,
            "State": state,
        })

        # Filas de pasos: solo Step/Action/Expected, igual al modelo exportado de Azure.
        for step_index, step in enumerate(steps_for_export, start=1):
            azure_rows.append({
                "ID": "",
                "Work Item Type": "",
                "Title": "",
                "Description": "",
                "Test Step": step.get("Step #", step_index),
                "Step Action": safe_text(
                    step.get("Action"),
                    "Acción no definida",
                ),
                "Step Expected": safe_text(
                    step.get("Expected value"),
                    "Resultado esperado no definido",
                ),
                "Area Path": "",
                "IDPadre": "",
                "Tipo Origen Proyecto": "",
                "Tiempo Real": "",
                "Assigned To": "",
                "State": "",
            })

        # --------------------------------------------------------
        # MATRIZ QA — se conserva sin cambios de columnas.
        # --------------------------------------------------------
        matriz_rows.append({
            "TestCaseId": case_id,
            "Title": title,
            "Requirement / Use Case": safe_text(
                coverage.get(
                    "Requirement / Use Case",
                    tc.get("Related Use Case")
                )
            ),
            "Criterion": safe_text(
                coverage.get("Criterion", tc.get("Criterion"))
            ),
            "Scenario": scenario,
            "Scenario Type": safe_text(
                tc.get("Scenario Type"),
                "No definido"
            ),
            "Description": description,
            "Preconditions": preconditions,
            "Validation Method": validation_method,
            "Coverage": coverage_value,
            "Alerts": alerts,
            "Effort": safe_text(
                tc.get("Effort"),
                "No definido"
            ),
        })

    df_azure = pd.DataFrame(azure_rows, columns=AZURE_COLUMNS)
    df_matriz = pd.DataFrame(matriz_rows, columns=MATRIZ_COLUMNS)

    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        # Hoja compatible con Azure DevOps Test Plans.
        df_azure.to_excel(
            writer,
            sheet_name="Azure Import",
            index=False,
        )

        # Matriz aprobada.
        df_matriz.to_excel(
            writer,
            sheet_name="Matriz QA",
            index=False,
        )

        for ws in writer.book.worksheets:
            ws.freeze_panes = "A2"
            ws.auto_filter.ref = ws.dimensions

            for column_cells in ws.columns:
                letter = column_cells[0].column_letter
                max_len = max(
                    len(str(cell.value or "")) for cell in column_cells
                )
                ws.column_dimensions[letter].width = min(
                    max(max_len + 2, 12), 60
                )

    output.seek(0)
    return output.getvalue()


# ============================================================
# PDF
# ============================================================
def pdf_text(value):
    return escape(safe_text(value)).replace("\n", "<br/>")


def _unique_values(cases, key):
    values = []
    seen = set()
    for tc in cases:
        value = re.sub(r"\s+", " ", safe_text(tc.get(key))).strip()
        if value and value.lower() not in seen:
            seen.add(value.lower())
            values.append(value)
    return values


def _pdf_bullets(story, items, style, prefix="•"):
    for item in items:
        text = safe_text(item)
        if text:
            story.append(Paragraph(f"{prefix} {pdf_text(text)}", style))


def _coverage_summary(data):
    rows = []
    for item in data.get("COVERAGE", []) or []:
        if not isinstance(item, dict):
            continue
        rows.append({
            "Requirement / Use Case": safe_text(item.get("Requirement / Use Case")),
            "Criterion": safe_text(item.get("Criterion")),
            "Scenario": safe_text(item.get("Scenario")),
            "Test Case": safe_text(item.get("Test Case")),
            "Validation Method": normalize_validation_method(
                item.get("Validation Method")
            ),
            "Coverage": normalize_coverage(item.get("Coverage")),
            "Alerts": safe_text(item.get("Alerts")),
        })
    return rows


def _register_reference_fonts():
    """Use Aptos from the supplied client reference when bundled."""
    font_dir = Path(__file__).resolve().parent / "fonts"
    regular = font_dir / "Aptos.ttf"
    bold = font_dir / "Aptos-Bold.ttf"
    italic = font_dir / "Aptos-BoldItalic.ttf"

    if regular.exists() and bold.exists():
        try:
            pdfmetrics.registerFont(TTFont("Aptos", str(regular)))
            pdfmetrics.registerFont(TTFont("Aptos-Bold", str(bold)))
            if italic.exists():
                pdfmetrics.registerFont(TTFont("Aptos-BoldItalic", str(italic)))
            return "Aptos", "Aptos-Bold", "Aptos-BoldItalic"
        except Exception:
            pass
    return "Helvetica", "Helvetica-Bold", "Helvetica-BoldOblique"


def _case_bullets_as_paragraphs(story, text, style):
    """Render source-style bullets without introducing a new table/column."""
    value = safe_text(text)
    if not value:
        return
    lines = [x.strip() for x in value.splitlines() if x.strip()]
    for line in lines:
        clean = re.sub(r"^[•\-]\s*", "", line)
        story.append(Paragraph(f"• {pdf_text(clean)}", style))


def create_pdf(data, config_key, source_name=""):
    """
    V16 — PDF alineado visual y estructuralmente con el Test Plan del cliente.

    Se conserva la estructura de referencia:
      1) bloque inicial tipo tabla: Descripción del Software, Objetivos,
         Elementos requeridos, Lista de ítems que no serán probados y Entregables;
      2) Test plan Ejecución;
      3) DETALLE DE LOS CASOS DE PRUEBA;
      4) por caso: Test case, SUMMARY, Producto, Módulo, Descripción,
         Resultado esperado de la prueba, Precondiciones y Caso de uso relacionado.

    No agrega columnas ni bloques de resumen que no existen en el documento de
    referencia. Los datos del caso siguen proviniendo del resultado QA.
    """
    buffer = io.BytesIO()
    regular_font, bold_font, italic_font = _register_reference_fonts()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=40,
        leftMargin=40,
        topMargin=36,
        bottomMargin=36,
        title="Test plan Ejecución — VERSION PREVIA — DRAFT",
        author="Agente QA",
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ref_title", parent=styles["Normal"], fontName=bold_font,
        fontSize=9.5, leading=11, spaceAfter=4,
    )
    cell_label = ParagraphStyle(
        "ref_label", parent=styles["Normal"], fontName=bold_font,
        fontSize=8, leading=9.5,
    )
    cell_body = ParagraphStyle(
        "ref_body", parent=styles["Normal"], fontName=regular_font,
        fontSize=8, leading=10,
    )
    cell_body_bold = ParagraphStyle(
        "ref_body_bold", parent=cell_body, fontName=bold_font,
    )
    section = ParagraphStyle(
        "ref_section", parent=styles["Normal"], fontName=bold_font,
        fontSize=9.5, leading=11, spaceBefore=8, spaceAfter=5,
    )
    case_head = ParagraphStyle(
        "ref_case", parent=styles["Normal"], fontName=bold_font,
        fontSize=8.8, leading=10.5, spaceBefore=6, spaceAfter=5,
    )
    summary = ParagraphStyle(
        "ref_summary", parent=styles["Normal"], fontName=bold_font,
        fontSize=9, leading=11, spaceBefore=4, spaceAfter=5,
    )
    body = ParagraphStyle(
        "ref_body_main", parent=styles["Normal"], fontName=regular_font,
        fontSize=8.3, leading=11, spaceAfter=4,
    )
    body_bold = ParagraphStyle(
        "ref_body_bold_main", parent=body, fontName=bold_font,
    )
    small_note = ParagraphStyle(
        "ref_note", parent=body, fontName=italic_font,
        fontSize=7.5, leading=9.5,
    )

    cases = data.get("TEST_CASES", []) or []
    products = _unique_values(cases, "Product")
    modules = _unique_values(cases, "Module")
    product_text = ", ".join(products) if products else "No definido en la fuente"
    module_text = ", ".join(modules) if modules else "No definido en la fuente"

    # Fuente del documento: solo se usa para identificar el entregable, no para inventar un ID.
    plan_name = safe_text(source_name, "Documentación proporcionada")

    # Descripción y objetivos se derivan de los casos generados, sin inventar información adicional.
    descriptions = []
    objectives = []
    for tc in cases:
        desc = safe_text(tc.get("Description"))
        if desc and desc.lower() not in {x.lower() for x in descriptions}:
            descriptions.append(desc)
        scenario = safe_text(tc.get("Scenario")) or build_case_title(
            tc, safe_text(tc.get("ID"), "CP-00001")
        )
        if scenario and scenario.lower() not in {x.lower() for x in objectives}:
            objectives.append(scenario)

    description_text = " ".join(descriptions[:3]) if descriptions else (
        "La documentación analizada contiene la información funcional utilizada "
        "para derivar los casos de prueba."
    )

    # ========================================================
    # BLOQUE INICIAL — MISMA ESTRUCTURA DE LA REFERENCIA
    # ========================================================
    story = []
    story.append(Paragraph(f"Test plan: {pdf_text(plan_name)}", title_style))

    objectives_flow = []
    for i, objective in enumerate(objectives[:10], start=1):
        objectives_flow.append(Paragraph(f"{i}. {pdf_text(objective)}", cell_body))

    elements_flow = [
        Paragraph("<i>Test plan para la documentación de los casos de prueba y su ejecución:</i>", cell_body),
        Spacer(1, 3),
        Paragraph("1. Software Screen Recorder o capturador de pantallas para grabar las evidencias.", cell_body),
        Spacer(1, 2),
        Paragraph("2. Tipos de pruebas:", cell_body),
        Paragraph("a. Pruebas de integración: Validar de manera general el funcionamiento de los módulos afectados por el ajuste.", cell_body),
        Paragraph("b. Pruebas funcionales: Validaciones de los módulos involucrados en el ajuste según casos de prueba diseñados.", cell_body),
    ]

    out_scope = [Paragraph("1. Cualquier otra funcionalidad no especificada en este documento.", cell_body)]
    deliverables = [
        Paragraph("1. Plan de pruebas.", cell_body),
        Paragraph("2. Informe de la ejecución de las pruebas.", cell_body),
        Paragraph("3. Archivo con evidencias de las pruebas realizadas para el proyecto.", cell_body),
        Paragraph("4. Requerimiento (opcional)", cell_body),
    ]

    # Build nested content as one cell to reproduce the client's two-column table.
    desc_cell = Paragraph(pdf_text(description_text), cell_body)
    obj_cell = objectives_flow or [Paragraph("No se identificaron objetivos explícitos en la información procesada.", cell_body)]
    elem_cell = elements_flow
    scope_cell = out_scope
    deliver_cell = deliverables

    table_data = [
        [Paragraph("Descripción del<br/>Software", cell_label), desc_cell],
        [Paragraph("Objetivos de las<br/>pruebas", cell_label), obj_cell],
        [Paragraph("Elementos<br/>requeridos", cell_label), elem_cell],
        [Paragraph("Lista de ítems<br/>que no serán<br/>probados", cell_label), scope_cell],
        [Paragraph("Entregables", cell_label), deliver_cell],
    ]

    plan_table = Table(table_data, colWidths=[105, 435])
    plan_table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.65, colors.HexColor("#777777")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(plan_table)
    story.append(Spacer(1, 18))
    story.append(Paragraph("Test plan Ejecución", section))
    story.append(HRFlowable(width="100%", thickness=0.65, color=colors.HexColor("#777777")))
    story.append(Spacer(1, 7))
    story.append(Paragraph("DETALLE DE LOS CASOS DE PRUEBA", section))

    # ========================================================
    # CASOS — MISMA SECUENCIA DE CAMPOS DE LA REFERENCIA
    # ========================================================
    for idx, tc in enumerate(cases, start=1):
        case_id = safe_text(tc.get("ID"), f"CP-{idx:05d}")
        title_value = build_case_title(tc, case_id)

        story.append(
            Paragraph(
                f"Test case {pdf_text(case_id)}: {pdf_text(title_value)}",
                case_head,
            )
        )
        story.append(Paragraph("SUMMARY", summary))
        story.append(Paragraph(
            f"<b>Producto:</b> {pdf_text(safe_text(tc.get('Product'), product_text))}",
            body,
        ))
        story.append(Paragraph(
            f"<b>Módulo:</b> {pdf_text(safe_text(tc.get('Module'), module_text))}",
            body,
        ))
        story.append(Paragraph(
            f"<b>Descripción:</b> {pdf_text(safe_text(tc.get('Description'), 'No definida en la fuente.'))}",
            body,
        ))

        expected_result = safe_text(tc.get("Expected Result"))
        story.append(Paragraph(
            f"<b>Resultado esperado de la prueba:</b> {pdf_text(expected_result or 'No definido en la fuente.')}",
            body,
        ))

        preconditions = safe_text(tc.get("Preconditions"))
        if preconditions:
            story.append(Paragraph("<b>Precondiciones:</b>", body))
            _case_bullets_as_paragraphs(story, preconditions, body)
        else:
            story.append(Paragraph(
                "<b>Precondiciones:</b> No se definieron precondiciones en la fuente.",
                body,
            ))

        related = safe_text(tc.get("Related Use Case"))
        if related:
            story.append(Paragraph("<b>Caso de uso relacionado:</b>", body))
            _case_bullets_as_paragraphs(story, related, body)
        else:
            story.append(Paragraph(
                "<b>Caso de uso relacionado:</b> No se identificó en la fuente.",
                body,
            ))

        # Los Steps se conservan, pero se presentan como texto corrido/numerado,
        # sin crear una tabla o columnas nuevas que no existen en la referencia.
        steps = safe_steps(tc)
        if steps:
            story.append(Spacer(1, 3))
            story.append(Paragraph("<b>Secuencia de prueba:</b>", body))
            for step in steps:
                num = safe_text(step.get("Step #"), "")
                action = safe_text(step.get("Action"))
                expected = safe_text(step.get("Expected value"))
                story.append(Paragraph(
                    f"{pdf_text(num)}. {pdf_text(action)}",
                    body,
                ))
                if expected:
                    story.append(Paragraph(
                        f"Resultado esperado: {pdf_text(expected)}",
                        body,
                    ))

        # Alerts are only printed when the case actually has one, avoiding a new
        # permanent section in every case.
        case_alerts = tc.get("Alerts", [])
        if isinstance(case_alerts, list) and case_alerts:
            for alert in case_alerts:
                if not isinstance(alert, dict):
                    continue
                alert_name = safe_text(alert.get("Alert"))
                reason = safe_text(alert.get("Reason"))
                validation = safe_text(alert.get("Validation Required"))
                note_text = " / ".join(x for x in [alert_name, reason, validation] if x)
                if note_text:
                    story.append(Paragraph(
                        f"<i>Alerta: {pdf_text(note_text)}</i>",
                        small_note,
                    ))

        if idx < len(cases):
            story.append(Spacer(1, 10))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()



def coverage_gate_or_stop(data):
    """Renderiza métricas y detiene el flujo si la cobertura es incompleta."""
    metrics = calculate_cu_coverage(
        data.get("TEST_CASES", []),
        data.get("USE_CASES", [])
    )
    render_cu_coverage(metrics)
    if not metrics["valid"]:
        st.error(
            "🚫 EXPORTACIÓN BLOQUEADA: cada Caso de Uso debe tener mínimo un Caso de Prueba."
        )
        st.stop()
    return metrics



# ============================================================
# REFERENCIA AZURE — SOLO LECTURA / PREVIEW
# ============================================================
_REFERENCE_LABELS = [
    "Producto:",
    "Módulo:",
    "Descripción:",
    "Resultado esperado de la prueba:",
    "Precondiciones:",
    "Caso de uso relacionado:",
]

def _reference_compare(detail):
    description = safe_text(detail.get("description"))
    steps = detail.get("steps") or []
    low = description.lower()
    return {
        "Producto dentro de Description": "producto:" in low,
        "Módulo dentro de Description": "módulo:" in low or "modulo:" in low,
        "Descripción dentro de Description": "descripción:" in low or "descripcion:" in low,
        "Resultado esperado dentro de Description": "resultado esperado de la prueba:" in low,
        "Precondiciones dentro de Description": "precondiciones:" in low,
        "Caso de uso relacionado dentro de Description": "caso de uso relacionado:" in low,
        "Steps con Action + Expected": bool(steps) and all(
            safe_text(s.get("Action")) and safe_text(s.get("Expected value"))
            for s in steps
        ),
    }

def _reference_description_pretty(description):
    raw = html.unescape(safe_text(description))
    raw = re.sub(r"(?i)<br\s*/?>", "\n", raw)
    raw = re.sub(r"(?i)</(div|p|li|ul|blockquote|strong|b)>", "\n", raw)
    raw = re.sub(r"(?i)<li[^>]*>", "• ", raw)
    raw = re.sub(r"<[^>]+>", "", raw)
    raw = html.unescape(raw)
    raw = raw.replace("\xa0", " ")
    raw = re.sub(r"[ \t]+", " ", raw)
    raw = re.sub(r"\n\s*\n\s*\n+", "\n\n", raw).strip()
    return raw

def _reference_description_sections(description):
    text = _reference_description_pretty(description)
    sections = []
    for idx, label in enumerate(_REFERENCE_LABELS):
        start = re.search(re.escape(label), text, flags=re.I)
        if not start:
            sections.append((label, "No encontrado en la referencia."))
            continue
        content_start = start.end()
        next_positions = []
        for next_label in _REFERENCE_LABELS[idx + 1:]:
            m = re.search(re.escape(next_label), text[content_start:], flags=re.I)
            if m:
                next_positions.append(content_start + m.start())
        end = min(next_positions) if next_positions else len(text)
        content = text[content_start:end].strip(" \n:-")
        sections.append((label, content or "Sin contenido visible."))
    return sections

def _build_reference_preview_case(reference_detail, generated_case):
    tc = dict(generated_case or {})
    reference_sections = dict(_reference_description_sections(reference_detail.get("description", "")))

    product = safe_text(tc.get("Product"), reference_sections.get("Producto:", ""), "Pendiente")
    module = safe_text(tc.get("Module"), reference_sections.get("Módulo:", ""), "Pendiente")
    description = safe_text(tc.get("Description"), tc.get("Scenario"), "Pendiente")
    expected = safe_text(tc.get("Expected Result"), "Pendiente")
    preconditions = safe_text(tc.get("Preconditions"), "Pendiente")
    related = safe_text(tc.get("Related Use Case"), "Pendiente")

    description_structured = build_azure_description(
        product=product,
        module=module,
        description=description,
        expected=expected,
        preconditions=preconditions,
        related_use_case=related,
    )

    return {
        "ID": safe_text(tc.get("ID")),
        "Title": build_case_title(tc, safe_text(tc.get("ID"), "CP-PREVIEW")),
        "Product": product,
        "Module": module,
        "Description": description_structured,
        "Preconditions": preconditions,
        "Expected Result": expected,
        "Related Use Case": related,
        "Steps": safe_steps(tc),
        "reference_test_case_id": reference_detail.get("id"),
        "reference_test_case_title": reference_detail.get("title"),
    }

# ============================================================
# INTERFAZ
# ============================================================
st.set_page_config(
    page_title=f"Agente QA {APP_VERSION}",
    layout="wide",
)

if "source_content" not in st.session_state:
    st.session_state.source_content = ""
if "source_name" not in st.session_state:
    st.session_state.source_name = ""
if "result_json" not in st.session_state:
    st.session_state.result_json = None
if "excel_data" not in st.session_state:
    st.session_state.excel_data = None
if "pdf_data" not in st.session_state:
    st.session_state.pdf_data = None

for _key, _default in {
    "azure_reference_plans": [],
    "azure_reference_plan_id": None,
    "azure_reference_suites": [],
    "azure_reference_suite_id": None,
    "azure_reference_cases": [],
    "azure_reference_case_id": None,
    "azure_reference_detail": None,
    "azure_reference_preview": None,
}.items():
    if _key not in st.session_state:
        st.session_state[_key] = _default

st.title(f"🤖 Agente QA {APP_VERSION} — Generador de Casos de Prueba")
st.caption(
    "VERSION PREVIA — DRAFT | PDF / DOCX / TXT / MD → análisis QA → Excel + PDF"
)

with st.sidebar:
    st.header("⚙️ Configuración")

    api_key = st.secrets.get(
        "GEMINI_API_KEY",
        os.getenv("GEMINI_API_KEY", ""),
    )

    if api_key:
        st.success("✅ GEMINI_API_KEY configurada")
    else:
        api_key = st.text_input(
            "🔑 Google Gemini API Key",
            type="password",
        )

    selected_model = st.selectbox(
        "Modelo",
        FALLBACK_MODELS,
        index=0,
    )

    selected_config = st.selectbox(
        "Formato de Excel",
        list(EXCEL_CONFIGS.keys()),
        index=0,
    )

    max_retries = st.number_input(
        "Máximo de reintentos",
        min_value=0,
        max_value=5,
        value=2,
    )

    wait_time = st.number_input(
        "Espera inicial (segundos)",
        min_value=1,
        max_value=60,
        value=10,
    )

    st.divider()
    st.subheader("🔐 Azure DevOps")
    st.caption(
        "Prueba de conexión en modo solo lectura. No crea, modifica ni elimina CP."
    )

    if st.button("🔌 Probar conexión con Azure DevOps", key="azure_test_connection"):
        try:
            with st.spinner("Verificando conexión con Azure DevOps..."):
                azure_result = test_connection()
            st.success("✅ Conexión con Azure DevOps correcta.")
            st.caption(
                f"Organización: {azure_result['organization']} | "
                f"Proyecto: {azure_result['project']} | "
                f"Work Item: {azure_result['work_item_type']}"
            )
            st.info(azure_result["message"])
        except AzureDevOpsError as exc:
            st.error(f"❌ No se pudo conectar con Azure DevOps: {exc}")
        except Exception as exc:
            st.error(f"❌ Error inesperado al probar Azure DevOps: {exc}")

    st.markdown("### 📋 Test Plans")
    st.caption(
        "Consulta de solo lectura. Solo se muestran los 10 Test Plans más recientes. "
        "No crea, edita ni elimina Test Plans, Suites ni Test Cases."
    )

    if st.button("📋 Consultar 10 Test Plans", key="azure_list_test_plans"):
        try:
            with st.spinner("Consultando los 10 Test Plans más recientes..."):
                plans_result = list_test_plans(limit=10)
            st.session_state.azure_reference_plans = plans_result.get("plans", [])
            st.session_state.azure_reference_plan_id = None
            st.session_state.azure_reference_suites = []
            st.session_state.azure_reference_suite_id = None
            st.session_state.azure_reference_cases = []
            st.session_state.azure_reference_case_id = None
            st.session_state.azure_reference_detail = None
            st.session_state.azure_reference_preview = None
            st.success(
                f"✅ Consulta correcta: {len(st.session_state.azure_reference_plans)} "
                "Test Plan(s) más recientes."
            )
            st.info(plans_result["message"])
        except AzureDevOpsError as exc:
            st.error(f"❌ No se pudieron consultar los Test Plans: {exc}")
        except Exception as exc:
            st.error(f"❌ Error inesperado al consultar Test Plans: {exc}")

    plans = st.session_state.get("azure_reference_plans", [])
    if plans:
        plan_options = [
            f"{_ui_text(p.get('id'), 'SIN ID')} — {_ui_text(p.get('name'), 'Test Plan sin nombre')}"
            for p in plans
        ]
        selected_plan_label = st.selectbox(
            "1️⃣ Test Plan",
            plan_options,
            key="azure_reference_plan_select",
        )
        selected_plan_id = plans[plan_options.index(selected_plan_label)].get("id")

        if st.button("🔎 Consultar Suites", key="azure_reference_get_suites"):
            try:
                with st.spinner("Consultando Suites del Test Plan seleccionado..."):
                    suites = list_test_suites(selected_plan_id)
                st.session_state.azure_reference_plan_id = selected_plan_id
                st.session_state.azure_reference_suites = suites
                st.session_state.azure_reference_suite_id = None
                st.session_state.azure_reference_cases = []
                st.session_state.azure_reference_case_id = None
                st.session_state.azure_reference_detail = None
                st.session_state.azure_reference_preview = None
                st.success(f"✅ {len(suites)} Suite(s) encontradas.")
            except AzureDevOpsError as exc:
                st.error(f"❌ No se pudieron consultar las Suites: {exc}")
            except Exception as exc:
                st.error(f"❌ Error inesperado al consultar Suites: {exc}")

    suites = st.session_state.get("azure_reference_suites", [])
    if suites:
        suite_options = [
            f"{_ui_text(s.get('id'), 'SIN ID')} — {_ui_text(s.get('name'), 'Suite sin nombre')}"
            for s in suites
        ]
        selected_suite_label = st.selectbox(
            "2️⃣ Suite",
            suite_options,
            key="azure_reference_suite_select",
        )
        selected_suite_id = suites[suite_options.index(selected_suite_label)].get("id")

        if st.button("🔎 Consultar Test Cases", key="azure_reference_get_cases"):
            try:
                plan_id_for_suite = st.session_state.get("azure_reference_plan_id") or selected_plan_id
                with st.spinner("Consultando Test Cases de la Suite seleccionada..."):
                    cases = list_test_cases(plan_id_for_suite, selected_suite_id)
                st.session_state.azure_reference_suite_id = selected_suite_id
                st.session_state.azure_reference_cases = cases
                st.session_state.azure_reference_case_id = None
                # Evita que Streamlit conserve la selección "None —" de una ejecución anterior.
                st.session_state.pop("azure_reference_case_select", None)
                st.session_state.azure_reference_detail = None
                st.session_state.azure_reference_preview = None
                st.success(f"✅ {len(cases)} Test Case(s) encontrados.")
            except AzureDevOpsError as exc:
                st.error(f"❌ No se pudieron consultar los Test Cases: {exc}")
            except Exception as exc:
                st.error(f"❌ Error inesperado al consultar Test Cases: {exc}")

    cases = st.session_state.get("azure_reference_cases", [])
    st.markdown("### 3️⃣ Test Case de referencia")
    if not cases and st.session_state.get("azure_reference_suite_id"):
        st.warning(
            "⚠️ La Suite seleccionada no tiene Test Cases consultables. "
            "Selecciona otra Suite del mismo Test Plan y vuelve a consultar."
        )
    if cases:
        # Solo llegan aquí Test Cases con ID válido; por seguridad filtramos cualquier residuo.
        valid_cases = [
            c for c in cases
            if c.get("id") is not None and str(c.get("id")).strip() != ""
        ]
        if not valid_cases:
            st.warning("⚠️ Azure devolvió registros, pero ninguno tiene un ID de Test Case válido.")
            st.stop()

        case_options = [
            f"{_ui_text(c.get('id'), 'SIN ID')} — {_ui_text(c.get('title'), 'Test Case sin título')}"
            for c in valid_cases
        ]
        selected_case_label = st.selectbox(
            "Selecciona un Test Case real de Azure para comparar su estructura",
            case_options,
            key="azure_reference_case_select",
        )
        selected_case_id = valid_cases[case_options.index(selected_case_label)].get("id")

        if st.button("🔬 Consultar y comparar", key="azure_reference_compare"):
            try:
                with st.spinner("Leyendo el Test Case real de Azure..."):
                    detail = get_test_case_detail(selected_case_id)
                if not selected_case_id:
                    raise AzureDevOpsError("El Test Case seleccionado no tiene un ID válido.")
                st.session_state.azure_reference_case_id = selected_case_id
                st.session_state.azure_reference_detail = detail
                st.session_state.azure_reference_preview = None
                st.success("✅ Comparación realizada. Solo se ejecutaron consultas GET.")
            except AzureDevOpsError as exc:
                st.error(f"❌ No se pudo consultar el Test Case de referencia: {exc}")
            except Exception as exc:
                st.error(f"❌ Error inesperado al consultar el Test Case: {exc}")

    reference_detail = st.session_state.get("azure_reference_detail")
    if reference_detail:
        st.markdown("## 📌 Test Case de referencia")
        st.markdown(
            f"**ID:** {safe_text(reference_detail.get('id'))}  \n"
            f"**Título:** {safe_text(reference_detail.get('title'))}  \n"
            f"**Estado:** {safe_text(reference_detail.get('state'))}  \n"
            f"**Area Path:** {safe_text(reference_detail.get('area_path'))}"
        )

        st.markdown("### 📝 Description real de Azure")
        for label, value in _reference_description_sections(reference_detail.get("description", "")):
            with st.container(border=True):
                st.markdown(f"**{label}**")
                st.write(value)

        st.markdown("### 🧪 Steps reales de Azure")
        steps_df = pd.DataFrame(reference_detail.get("steps") or [])
        if not steps_df.empty:
            display_cols = [c for c in ["Step #", "Action", "Expected value"] if c in steps_df.columns]
            st.dataframe(steps_df[display_cols], width="stretch", hide_index=True)
        else:
            st.warning("⚠️ El Test Case de referencia no tiene Steps legibles.")

        checks = _reference_compare(reference_detail)
        comparison_rows = [
            {"Elemento": label, "Cumple": "✅" if ok else "⚠️"}
            for label, ok in checks.items()
        ]
        st.markdown("### 🔎 Comparación contra nuestra estructura aprobada")
        st.dataframe(pd.DataFrame(comparison_rows), width="stretch", hide_index=True)

        if all(checks.values()):
            st.success(
                "✅ La referencia contiene la estructura aprobada. "
                "Ahora el agente puede preparar un CP nuevo en PREVIEW, sin enviarlo a Azure."
            )
        else:
            st.warning(
                "⚠️ La referencia no contiene todos los elementos esperados. "
                "Se conserva como referencia, pero no se inventan los elementos faltantes."
            )

        current_result = st.session_state.get("result_json")
        if current_result and all(checks.values()):
            st.markdown("### 4️⃣ Preparar nuevo Test Case — PREVIEW")
            st.caption(
                "El agente toma la estructura del Test Case de referencia y la aplica "
                "al CP generado. En esta etapa NO se crea ni modifica ningún recurso en Azure."
            )
            if st.button(
                "🧩 Preparar CP nuevo con la estructura de referencia",
                key="azure_prepare_new_cp_preview",
            ):
                try:
                    generated_cases = current_result.get("TEST_CASES", []) or []
                    if not generated_cases:
                        raise ValueError("No hay Test Cases generados para preparar el preview.")
                    st.session_state.azure_reference_preview = _build_reference_preview_case(
                        reference_detail,
                        generated_cases[0],
                    )
                except Exception as exc:
                    st.error(f"❌ No se pudo preparar el PREVIEW: {exc}")

        preview = st.session_state.get("azure_reference_preview")
        if preview:
            st.markdown("### 👀 CP nuevo antes de enviarlo a Azure")
            st.info(
                "PUNTO DE CONTROL: este CP es únicamente una vista previa. "
                "No se ha realizado ningún POST/creación en Azure DevOps."
            )
            st.markdown(f"**Title:** {safe_text(preview.get('Title'))}")

            st.markdown("#### Description")
            for label, value in _reference_description_sections(preview.get("Description", "")):
                with st.container(border=True):
                    st.markdown(f"**{label}**")
                    st.write(value)

            st.markdown("#### Steps")
            preview_steps = pd.DataFrame(preview.get("Steps") or [])
            if not preview_steps.empty:
                cols = [c for c in ["Step #", "Action", "Expected value"] if c in preview_steps.columns]
                st.dataframe(preview_steps[cols], width="stretch", hide_index=True)
            else:
                st.warning("⚠️ El CP nuevo no tiene Steps para revisar.")

            st.warning(
                "🛑 FIN DE ESTA FASE: el CP queda listo para revisión funcional. "
                "Todavía no existe ninguna acción de creación, actualización o eliminación en Azure."
            )

            # ============================================================
            # REVISION FUNCIONAL — SIN ESCRITURA EN AZURE
            # ============================================================
            st.markdown("### 👩‍💻 5️⃣ Revisión funcional")
            st.caption(
                "Esta etapa permite validar manualmente el CP preparado antes de cualquier futura integración de creación en Azure. "
                "No realiza POST, PATCH ni DELETE en Azure DevOps."
            )

            review_product = bool(safe_text(preview.get("Product")))
            review_module = bool(safe_text(preview.get("Module")))
            review_description = bool(safe_text(preview.get("Description")))
            review_expected = bool(safe_text(preview.get("Expected Result")))
            review_preconditions = bool(safe_text(preview.get("Preconditions")))
            review_related = bool(safe_text(preview.get("Related Use Case")))
            review_steps = bool(preview.get("Steps")) and all(
                safe_text(step.get("Action")) and safe_text(step.get("Expected value"))
                for step in (preview.get("Steps") or [])
            )

            review_rows = [
                {"Validación funcional": "Producto en Description", "Estado": "✅" if review_product else "⚠️"},
                {"Validación funcional": "Módulo en Description", "Estado": "✅" if review_module else "⚠️"},
                {"Validación funcional": "Descripción funcional completa", "Estado": "✅" if review_description else "⚠️"},
                {"Validación funcional": "Resultado esperado en Description", "Estado": "✅" if review_expected else "⚠️"},
                {"Validación funcional": "Precondiciones en Description", "Estado": "✅" if review_preconditions else "⚠️"},
                {"Validación funcional": "Caso de uso relacionado identificado", "Estado": "✅" if review_related else "⚠️"},
                {"Validación funcional": "Steps con Action + Expected", "Estado": "✅" if review_steps else "⚠️"},
            ]
            st.dataframe(pd.DataFrame(review_rows), width="stretch", hide_index=True)

            review_ok = all([
                review_product, review_module, review_description, review_expected,
                review_preconditions, review_related, review_steps,
            ])

            if not review_ok:
                st.error(
                    "❌ La revisión funcional no puede aprobarse todavía. "
                    "Hay elementos que requieren validación funcional."
                )
            else:
                st.success(
                    "✅ El CP contiene la estructura mínima preparada para revisión funcional. "
                    "La aprobación sigue siendo manual y no crea el Test Case en Azure."
                )

            review_decision = st.radio(
                "Decisión de la revisión funcional",
                ["Pendiente", "Aprobado funcionalmente", "Requiere ajustes"],
                key="azure_functional_review_decision",
                horizontal=True,
            )

            if review_decision == "Aprobado funcionalmente" and review_ok:
                st.success(
                    "🟢 Revisión funcional aprobada. El CP queda preparado como siguiente insumo. "
                    "No se ha enviado ni creado en Azure DevOps."
                )
            elif review_decision == "Requiere ajustes":
                st.warning(
                    "🟡 El CP requiere ajustes funcionales. Se mantiene en PREVIEW y no se realiza ninguna escritura en Azure."
                )



st.subheader("📁 Carga de Documento")

st.info(
    "Formatos de HU: TXT, MD, PDF, DOCX, XLSX/CSV. "
    "Para PDF escaneado se requiere OCR; esta versión no inventa "
    "texto que no pueda extraer."
)

uploaded = st.file_uploader(
    "Arrastra o selecciona un documento",
    type=["txt", "md", "pdf", "docx", "xlsx", "xls", "csv"],
)

source_text = st.session_state.source_content

if uploaded:
    # Procesar solo cuando cambia el archivo.
    if st.session_state.source_name != uploaded.name:
        try:
            with st.spinner(f"Procesando {uploaded.name}..."):
                source_text = extract_source(uploaded)

            st.session_state.source_content = source_text
            st.session_state.source_name = uploaded.name
            st.session_state.result_json = None
            st.session_state.excel_data = None
            st.session_state.pdf_data = None

            st.success(f"✅ {uploaded.name} procesado correctamente.")

        except Exception as exc:
            st.session_state.source_content = ""
            st.session_state.source_name = ""
            st.session_state.result_json = None
            st.session_state.excel_data = None
            st.session_state.pdf_data = None

            st.error(f"❌ No se pudo procesar el archivo: {exc}")

if source_text:
    with st.expander("📄 Vista previa del contenido", expanded=True):
        st.text_area(
            "Contenido",
            source_text[:5000],
            height=250,
            disabled=True,
        )

        c1, c2, c3 = st.columns(3)
        c1.metric("Caracteres", len(source_text))
        c2.metric("Líneas", len(source_text.splitlines()))
        c3.metric("Palabras", len(source_text.split()))

else:
    source_text = st.text_area(
        "✏️ O ingresa el texto manualmente",
        height=220,
        placeholder=(
            "Pega aquí la Historia de Usuario o documentación fuente..."
        ),
    )

    if source_text:
        st.session_state.source_content = source_text


st.divider()
st.subheader("🧪 Generación QA")

if st.button(
    "🚀 Generar casos de prueba",
    type="primary",
    disabled=not bool(source_text.strip()),
):
    try:
        with st.spinner(
            "Analizando documentación y generando casos..."
        ):
            result = generate_qa_data(
                load_prompt(),
                source_text,
                api_key,
                selected_model,
                0.0,
                int(max_retries),
                int(wait_time),
            )

        # Regla obligatoria: validar cobertura antes de exportar.
        coverage_metrics = coverage_gate_or_stop(result)

        st.session_state.result_json = result
        st.session_state.excel_data = create_excel(
            result,
            selected_config,
        )
        st.session_state.pdf_data = create_pdf(
            result,
            selected_config,
            st.session_state.get("source_name", ""),
        )

        st.success(
            f"✅ Generación completada: "
            f"{len(result['TEST_CASES'])} casos."
        )

    except Exception as exc:
        st.error(f"❌ Error durante la generación: {exc}")


result = st.session_state.result_json

if result:
    st.divider()
    st.subheader("📊 Resultados")

    c1, c2, c3 = st.columns(3)
    c1.metric("Casos", len(result.get("TEST_CASES", [])))
    c2.metric("Alertas", len(result.get("ALERTS", [])))
    c3.metric("CUs", len(result.get("USE_CASES", [])))

    current_coverage = calculate_cu_coverage(
        result.get("TEST_CASES", []),
        result.get("USE_CASES", [])
    )
    render_cu_coverage(current_coverage)

    # ========================================================
    # EDITOR DE CASOS DE PRUEBA — EXPERIENCIA TIPO AZURE
    # V13: un CU por Test Case + eliminar CP
    # ========================================================
    st.subheader("✏️ Editar caso de prueba")
    st.caption(
        "Revisa y ajusta el caso antes de descargar Excel/PDF. "
        "Cada Test Case debe conservar un único Caso de Uso relacionado."
    )

    case_options = []
    for idx, tc in enumerate(result.get("TEST_CASES", [])):
        case_id = safe_text(tc.get("ID"), f"CASO-{idx + 1:05d}")
        case_title = build_case_title(tc, case_id)
        case_options.append(f"{case_id} — {case_title[:100]}")

    if case_options:
        selected_case_label = st.selectbox(
            "Selecciona el caso que deseas editar",
            case_options,
            key="qa_editor_selected_case",
        )
        selected_index = case_options.index(selected_case_label)
        selected_case = result["TEST_CASES"][selected_index]
        selected_case_id = safe_text(
            selected_case.get("ID"),
            f"CASO-{selected_index + 1:05d}",
        )

        st.markdown(f"### {selected_case_id}")
        st.info(
            "El ID no se puede modificar. La relación funcional del caso "
            "debe corresponder a un solo Caso de Uso."
        )

        editor_result = render_azure_style_editor(selected_case, selected_index)

        if editor_result == "saved":
            try:
                coverage_after_edit = calculate_cu_coverage(
                    result.get("TEST_CASES", []),
                    result.get("USE_CASES", [])
                )
                if not coverage_after_edit["valid"]:
                    st.error(
                        "🚫 El cambio dejaría un CU sin CP. "
                        "Corrige la relación antes de exportar."
                    )
                else:
                    st.session_state.excel_data = create_excel(result, selected_config)
                    st.session_state.pdf_data = create_pdf(
                        result,
                        selected_config,
                        st.session_state.get("source_name", ""),
                    )
                    st.success(
                        f"✅ {selected_case_id} actualizado. "
                        "Excel y PDF regenerados."
                    )
                    st.rerun()
            except Exception as exc:
                st.error(f"❌ No se pudo guardar el cambio: {exc}")

        st.divider()
        st.markdown("### 🗑️ Eliminar caso de prueba")

        st.warning(
            "Esta acción elimina el Test Case seleccionado de la generación actual "
            "y también elimina su relación en COVERAGE. No elimina otros casos."
        )

        confirm_delete = st.checkbox(
            f"Confirmo que quiero eliminar {selected_case_id}",
            key=f"v13_confirm_delete_{selected_index}",
        )

        if st.button(
            "🗑️ Eliminar CP seleccionado",
            type="secondary",
            disabled=not confirm_delete,
            key=f"v13_delete_cp_{selected_index}",
        ):
            deleted_id = selected_case_id
            # Simular eliminación antes de aplicarla para preservar 1 CP mínimo por CU.
            candidate_cases = [
                tc for i, tc in enumerate(result.get("TEST_CASES", []))
                if i != selected_index
            ]
            deletion_coverage = calculate_cu_coverage(
                candidate_cases,
                result.get("USE_CASES", [])
            )

            if not deletion_coverage["valid"]:
                st.error(
                    "🚫 No se puede eliminar este CP porque dejaría al menos un CU sin cobertura."
                )
            elif delete_test_case(result, selected_index):
                st.session_state.excel_data = create_excel(result, selected_config)
                st.session_state.pdf_data = create_pdf(
                    result,
                    selected_config,
                    st.session_state.get("source_name", ""),
                )
                st.success(f"✅ {deleted_id} eliminado.")
                st.rerun()
    else:
        st.info("No hay Test Cases para editar.")

    if not safe_text(EXCEL_CONFIGS[selected_config].get("area_path")) or not safe_text(EXCEL_CONFIGS[selected_config].get("assigned_to")):
        st.warning(
            "⚠️ El Excel conserva la estructura Azure y agrega Tipo Origen Proyecto = Proyecto. "
            "Antes de importar, verifica Assigned To en EXCEL_CONFIGS "
            "con valores reales de tu proyecto/organización; no se inventan automáticamente."
        )

    coverage_ok_for_download = calculate_cu_coverage(
        result.get("TEST_CASES", []),
        result.get("USE_CASES", [])
    )["valid"]

    if not coverage_ok_for_download:
        st.error(
            "🚫 Descargas deshabilitadas: la cobertura mínima por CU no se cumple."
        )

    # IMPORTANTE: el Excel se reconstruye SIEMPRE con el estado actual de
    # `result`, después de aplicar cualquier edición realizada en el editor.
    # Así, cambios en Description, Preconditions, Steps, Expected, Related CU,
    # etc. quedan reflejados en el archivo descargado. No se reutiliza una
    # copia anterior de session_state.excel_data que pudiera quedar desactualizada.
    current_excel_data = create_excel(result, selected_config)
    st.session_state.excel_data = current_excel_data

    st.download_button(
        "📊 Descargar Excel",
        data=current_excel_data,
        file_name=(
            f"QA_DRAFT_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
        ),
        mime=(
            "application/vnd.openxmlformats-officedocument."
            "spreadsheetml.sheet"
        ),
        disabled=not coverage_ok_for_download,
    )

    st.download_button(
        "📄 Descargar PDF",
        data=st.session_state.pdf_data,
        file_name=(
            f"QA_DRAFT_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf"
        ),
        mime="application/pdf",
        disabled=not coverage_ok_for_download,
    )

    with st.expander("🔎 Ver JSON generado", expanded=False):
        st.json(result)

    st.subheader("🧪 Casos generados")

    preview_rows = []

    for tc in result["TEST_CASES"]:
        preview_rows.append({
            "ID": safe_text(tc.get("ID")),
            "Title": build_case_title(
                tc,
                normalize_case_id(
                    tc.get("ID"),
                    safe_text(tc.get("Module"), "GENERAL"),
                    len(preview_rows) + 1,
                    EXCEL_CONFIGS[selected_config]["title_prefix"],
                ),
            ),
            "Module": safe_text(tc.get("Module")),
            "Scenario Type": safe_text(tc.get("Scenario Type")),
            "Steps": len(safe_steps(tc)),
        })

    st.dataframe(
        pd.DataFrame(preview_rows),
        width="stretch",
    )
